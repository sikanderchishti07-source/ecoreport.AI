"""Soil and water monitoring report.

Deliberately short. The air and noise reports run long because a time
series needs its method, its capture rate, its meteorology and a chart per
pollutant explained. A laboratory report is a different document: the
finding is a table of numbers against limits, and everything else exists to
say where the samples came from and which column of the standard was used.

Five sections, two appendices, three to five pages before the certificates.

The one section that carries real weight is the basis of assessment. A
soil limit is chosen by particle size, land use and depth; get the land use
wrong and every number on the page is compared against the wrong column
while the report reads as entirely correct. So the context is stated in its
own table near the front, and Appendix A prints all five land-use columns
with the applied one marked — a reader can see in one glance what the same
results would have been judged against elsewhere.

Style follows noise_generate.py exactly: Times New Roman, body 13 pt,
headings 16/14/13, captions 13, the same navy, the same table treatment.
Two reports from one company should not look like two companies.
"""
from __future__ import annotations

import logging
import os
from datetime import datetime
from typing import Dict, List, Optional

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

import soil_water_limits as L
from sample_calc import land_use_comparison

log = logging.getLogger(__name__)

NAVY = RGBColor(0x0F, 0x3D, 0x6E)
NAVY_HEX = "0F3D6E"
GREEN_HEX = "1E7D4F"
RED_HEX = "B03A2E"
AMBER_HEX = "B06A00"
MUT = RGBColor(0x6B, 0x72, 0x80)
INK = RGBColor(0x20, 0x24, 0x2B)
RED = RGBColor(0xB0, 0x3A, 0x2E)
HAIR = "D9DDE3"
ROW = "F4F6F9"
EXCEED_FILL = "FBE9E7"

ASSETS = os.path.join(os.path.dirname(os.path.abspath(__file__)), "assets")

MEDIUM_TITLES = {
    "soil": "Soil Monitoring Report",
    "water": "Water Monitoring Report",
    "sediment": "Sediment Monitoring Report",
}

STANDARD_CITATIONS = {
    "ncec_soil": ("Appendix (1) of the Executive Regulation for the "
                  "Prevention and Remediation of Soil Pollution, issued "
                  "under Royal Decree No. (m/165) dated 19/11/1441H"),
    "ncec_water_ambient": ("Appendix (1) of the Executive Regulations for "
                           "the Protection of Aqueous Media from Pollution, "
                           "issued under Royal Decree No. (m/165) dated "
                           "19/11/1441H"),
    "ncec_water_discharge": ("Appendices (2) and (3) of the Executive "
                             "Regulations for the Protection of Aqueous "
                             "Media from Pollution, issued under Royal "
                             "Decree No. (m/165) dated 19/11/1441H"),
    "ads_81_2017": "Abu Dhabi specification ADS 81/2017",
}


# ---------------------------------------------------------------------------
# small helpers — same shapes as noise_generate.py
# ---------------------------------------------------------------------------
def _bleed(paragraph):
    """Turn the paragraph's inline picture into a page-anchored one at 0,0."""
    inline = paragraph._p.find(qn("w:r") + "/" + qn("w:drawing") + "/"
                               + qn("wp:inline"))
    if inline is None:
        drawings = paragraph._p.findall(".//" + qn("wp:inline"))
        if not drawings:
            return
        inline = drawings[0]
    anchor = OxmlElement("wp:anchor")
    for k, v in (("distT", "0"), ("distB", "0"), ("distL", "0"),
                 ("distR", "0"), ("simplePos", "0"), ("relativeHeight", "1"),
                 ("behindDoc", "1"), ("locked", "0"), ("layoutInCell", "1"),
                 ("allowOverlap", "1")):
        anchor.set(k, v)
    sp = OxmlElement("wp:simplePos")
    sp.set("x", "0")
    sp.set("y", "0")
    anchor.append(sp)
    for tag, rel in (("positionH", "page"), ("positionV", "page")):
        el = OxmlElement("wp:" + tag)
        el.set("relativeFrom", rel)
        off = OxmlElement("wp:posOffset")
        off.text = "0"
        el.append(off)
        anchor.append(el)
    for child in list(inline):
        tag = child.tag.split("}")[-1]
        if tag == "extent":
            anchor.append(child)
            eff = OxmlElement("wp:effectExtent")
            for e in ("l", "t", "r", "b"):
                eff.set(e, "0")
            anchor.append(eff)
            anchor.append(OxmlElement("wp:wrapNone"))
        elif tag != "effectExtent":
            anchor.append(child)
    drawing = inline.getparent()
    drawing.remove(inline)
    drawing.append(anchor)


def _tight(p, before=0, after=6):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    return p


def _table_width(table, mm: float):
    tblPr = table._tbl.tblPr
    for old in tblPr.findall(qn("w:tblW")):
        tblPr.remove(old)
    w = OxmlElement("w:tblW")
    w.set(qn("w:w"), str(int(mm * 56.7)))
    w.set(qn("w:type"), "dxa")
    tblPr.append(w)


def _shade(cell, hexval):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexval)
    tcPr.append(sh)


def _borders(cell, **spec):
    tcPr = cell._tc.get_or_add_tcPr()
    tb = OxmlElement("w:tcBorders")
    for tag in ("top", "bottom", "left", "right"):
        el = OxmlElement("w:" + tag)
        v = spec.get(tag)
        if v is None:
            el.set(qn("w:val"), "nil")
        else:
            colour, sz = v
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), str(sz))
            el.set(qn("w:color"), colour)
        tb.append(el)
    tcPr.append(tb)


def _repeat_header(row):
    """Mark a table row as a header that repeats on every page.

    A twenty-four row results table crosses a page break, and the second
    page without column headings is a grid of numbers nobody can read.
    """
    trPr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    trPr.append(el)


def _cell(cell, text, size=11, bold=False, colour=INK, align="left",
          italic=False):
    p = cell.paragraphs[0]
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT,
                   "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    _tight(p, 2, 2)
    r = p.add_run(str(text))
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.name = "Times New Roman"
    r.font.color.rgb = colour
    return p


def _table(doc, headers, rows, widths_mm, aligns=None, size=None,
           fills=None, bolds=None):
    """Same treatment as the noise report's tables.

    `fills` and `bolds` are parallel grids of the same shape as `rows`, used
    to shade an exceeding cell rather than relying on the reader to compare
    every number against the limit column themselves.
    """
    if size is None:
        n = len(headers)
        size = 9 if n >= 9 else (9.5 if n >= 8 else (10.5 if n >= 6 else 11))
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for j, wmm in enumerate(widths_mm):
        for c in t.columns[j].cells:
            c.width = Mm(wmm)
    for j, h in enumerate(headers):
        c = t.cell(0, j)
        _shade(c, NAVY_HEX)
        _cell(c, h, size=size, bold=True, colour=RGBColor(255, 255, 255),
              align="center")
        _borders(c, top=(NAVY_HEX, 4), bottom=(NAVY_HEX, 4))
    _repeat_header(t.rows[0])
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            c = t.cell(i + 1, j)
            fill = (fills[i][j] if fills else None)
            if fill:
                _shade(c, fill)
            elif i % 2 == 1:
                _shade(c, ROW)
            _cell(c, v, size=size,
                  bold=bool(bolds[i][j]) if bolds else False,
                  colour=RED if fill == EXCEED_FILL else INK,
                  align=(aligns[j] if aligns else "center"))
            _borders(c, bottom=(HAIR, 2))
    return t


def _group_row(table, index, label, span, size=9.5):
    """Merge a results row into a single shaded band naming the group."""
    row = table.rows[index]
    merged = row.cells[0]
    for k in range(1, span):
        merged = merged.merge(row.cells[k])
    for p in list(merged.paragraphs)[1:]:
        p._p.getparent().remove(p._p)
    merged.paragraphs[0].text = ""
    _shade(merged, "E8EDF3")
    _cell(merged, label, size=size, bold=True, colour=NAVY, align="left")
    _borders(merged, bottom=(HAIR, 2))


def _callout(doc, text, colour_hex, fill_hex, mark):
    t = doc.add_table(rows=1, cols=1)
    t.autofit = False
    body = t.cell(0, 0)
    body.width = Mm(164)
    _table_width(t, 164)
    _shade(body, fill_hex)
    tcPr = body._tc.get_or_add_tcPr()
    tb = OxmlElement("w:tcBorders")
    for tag in ("top", "bottom", "right"):
        el = OxmlElement("w:" + tag)
        el.set(qn("w:val"), "nil")
        tb.append(el)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:color"), colour_hex)
    tb.append(left)
    tcPr.append(tb)
    p = _tight(body.paragraphs[0], 3, 3)
    r = p.add_run(f"{mark}  ")
    r.font.size = Pt(9)
    r.font.bold = True
    r.font.color.rgb = RGBColor(int(colour_hex[0:2], 16),
                                int(colour_hex[2:4], 16),
                                int(colour_hex[4:6], 16))
    r2 = p.add_run(text)
    r2.font.size = Pt(9)
    _tight(doc.add_paragraph(), 0, 2)


def _heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    if level == 1:
        pPr = h._p.get_or_add_pPr()
        b = OxmlElement("w:pBdr")
        bt = OxmlElement("w:bottom")
        bt.set(qn("w:val"), "single")
        bt.set(qn("w:sz"), "8")
        bt.set(qn("w:color"), NAVY_HEX)
        b.append(bt)
        pPr.append(b)
    return h


def _caption(doc, text):
    p = doc.add_paragraph(text, style="Caption")
    _tight(p, 4, 8)
    for r in p.runs:
        r.font.bold = False
        r.font.size = Pt(13)
        r.font.color.rgb = NAVY
    return p


def _muted(doc, text, size=10):
    p = _tight(doc.add_paragraph(), 0, 8)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.italic = True
    r.font.color.rgb = MUT
    return p


def _bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(it, style="List Bullet")
        _tight(p, 0, 2)
        for r in p.runs:
            r.font.size = Pt(13)
            r.font.name = "Times New Roman"
    return doc


def _body(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def _toc_anchor(doc, instr):
    """An empty TOC field for report.fields.build_indexes to fill."""
    p = _tight(doc.add_paragraph(), 0, 0)
    for kind, txt in (("begin", None), (None, instr), ("separate", None),
                      (None, None), ("end", None)):
        r = OxmlElement("w:r")
        if kind:
            f = OxmlElement("w:fldChar")
            f.set(qn("w:fldCharType"), kind)
            r.append(f)
        elif txt:
            it = OxmlElement("w:instrText")
            it.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            it.text = txt
            r.append(it)
        else:
            t = OxmlElement("w:t")
            t.text = ""
            r.append(t)
        p._p.append(r)
    return p


def _fmt_dt(dt: Optional[datetime]) -> str:
    if not dt:
        return "\u2014"
    return dt.strftime("%d %B %Y %H:%M")


def _date_only(dt: Optional[datetime]) -> str:
    if not dt:
        return "\u2014"
    return dt.strftime("%d %B %Y")


def _coord(v) -> str:
    try:
        return f"{float(v):.5f}"
    except (TypeError, ValueError):
        return "\u2014"


# ---------------------------------------------------------------------------
# the report
# ---------------------------------------------------------------------------
def generate_sample_report(campaign, settings, samples, summary,
                           figs: Dict[str, str], out_path: str,
                           site_map_path: Optional[str] = None,
                           site_photo_paths: Optional[List[str]] = None,
                           cover_photo_path: Optional[str] = None,
                           calibration_items: Optional[List[dict]] = None,
                           license_image_paths: Optional[List[str]] = None,
                           work_dir: Optional[str] = None) -> str:
    medium = settings.medium or "soil"
    title = MEDIUM_TITLES.get(medium, "Monitoring Report")
    standard = summary.standard or "none"
    citation = STANDARD_CITATIONS.get(standard)

    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(13)
    st.paragraph_format.space_after = Pt(6)
    st.paragraph_format.line_spacing = 1.0
    for lvl, sz in (("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 13)):
        h = doc.styles[lvl]
        h.font.name = "Times New Roman"
        h.font.size = Pt(sz)
        h.font.bold = True
        h.font.color.rgb = NAVY
    cap = doc.styles["Caption"]
    cap.font.name = "Times New Roman"
    cap.font.color.rgb = NAVY

    _tn = {"n": 0}

    def tnum() -> int:
        _tn["n"] += 1
        return _tn["n"]

    _fg = {"n": 0}

    def fnum() -> int:
        _fg["n"] += 1
        return _fg["n"]

    # Appendix letters settled before the body is written, because section 4
    # cites the certificate appendix by letter long before it is reached.
    _appendix: Dict[str, str] = {}
    _next = 0
    show_land_use_table = (standard == "ncec_soil"
                           and settings.default_context.particle_size
                           in L.PARTICLE_SIZES)
    if show_land_use_table:
        _appendix["limits"] = chr(ord("A") + _next)
        _next += 1
    _appendix["certificates"] = chr(ord("A") + _next)
    _next += 1
    if license_image_paths:
        _appendix["licence"] = chr(ord("A") + _next)
        _next += 1

    from report.imaging import slim
    wd = work_dir or os.path.dirname(os.path.abspath(out_path))

    def pic(p, path, width_mm):
        p.add_run().add_picture(slim(path, wd) or path, width=Mm(width_mm))

    # ---- cover -----------------------------------------------------------
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Mm(210), Mm(297)
    sec.top_margin = sec.bottom_margin = Mm(0)
    sec.left_margin = sec.right_margin = Mm(0)
    sec.header_distance = sec.footer_distance = Mm(0)
    sec.gutter = Mm(0)

    logo = os.path.join(ASSETS, "logo_left.png")
    plate_done = False
    try:
        from report.cover_plate import date_range, render_cover
        plate = os.path.join(wd, "sample_cover_plate.png")
        render_cover(
            plate,
            survey_dates=[date_range(campaign.monitoring_start,
                                     campaign.monitoring_end),
                          date_range(campaign.monitoring_start,
                                     campaign.monitoring_end, abbrev=True)],
            location=(campaign.site_name or campaign.project_name or "\u2014"),
            client=campaign.client or "\u2014",
            report_number=campaign.report_number or "\u2014",
            revision=campaign.revision or "00",
            issue_date=_date_only(getattr(campaign, "reporting_date", None)),
            photo_path=cover_photo_path,
            logo_path=logo if os.path.exists(logo) else None)
        p = _tight(doc.add_paragraph(), 0, 0)
        p.add_run().add_picture(plate, width=Mm(210), height=Mm(297))
        _bleed(p)
        for r in p.runs:
            r.font.size = Pt(1)
        plate_done = True
    except Exception:  # noqa: BLE001
        log.warning("cover plate failed — falling back to a plain cover",
                    exc_info=True)

    if not plate_done:
        for _ in range(6):
            _tight(doc.add_paragraph(), 0, 0)
        for text, size, bold, colour in (
                (title.upper(), 26, True, NAVY),
                (campaign.project_name or "", 15, True, NAVY),
                (f"{campaign.site_name or ''}  \u00b7  "
                 f"{len(samples)} sample(s)", 10, False, MUT)):
            p = _tight(doc.add_paragraph(), 0, 6)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text)
            r.font.size = Pt(size)
            r.font.bold = bold
            r.font.color.rgb = colour
        info = [("Client", campaign.client or "\u2014"),
                ("Site", campaign.site_name or "\u2014"),
                ("Sampling date", _date_only(campaign.monitoring_start)),
                ("Report number", campaign.report_number or "\u2014"),
                ("Revision and issue date",
                 f"{campaign.revision or '00'}    \u00b7    "
                 f"{_date_only(getattr(campaign, 'reporting_date', None))}")]
        t = doc.add_table(rows=len(info), cols=2)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.autofit = False
        _table_width(t, 180)
        for i, (lab, val) in enumerate(info):
            a, b = t.cell(i, 0), t.cell(i, 1)
            a.width = Mm(52)
            b.width = Mm(128)
            _cell(a, lab, size=8.5, colour=MUT)
            _cell(b, val, size=10, bold=True)
            for c in (a, b):
                _borders(c, bottom=(HAIR, 2))

    # ---- body section ----------------------------------------------------
    body = doc.add_section(WD_SECTION.NEW_PAGE)
    body.top_margin = body.bottom_margin = Mm(22)
    body.left_margin = body.right_margin = Mm(23)

    hdr = body.header
    hdr.is_linked_to_previous = False
    ht = hdr.add_table(rows=1, cols=2, width=Mm(164))
    ht.autofit = False
    hc1, hc2 = ht.cell(0, 0), ht.cell(0, 1)
    hc1.width = Mm(50)
    hc2.width = Mm(114)
    hp = _tight(hc1.paragraphs[0], 0, 0)
    if os.path.exists(logo):
        hp.add_run().add_picture(logo, width=Mm(22))
    _cell(hc2, title, size=10, bold=True, colour=NAVY, align="right")
    p2 = _tight(hc2.add_paragraph(), 0, 0)
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p2.add_run("  \u00b7  ".join(
        x for x in (campaign.project_name, campaign.report_number,
                    f"Rev {campaign.revision or '00'}") if x))
    r.font.size = Pt(7.5)
    r.font.color.rgb = MUT
    for c in (hc1, hc2):
        _borders(c, bottom=(NAVY_HEX, 6))

    ftr = body.footer
    ftr.is_linked_to_previous = False
    fp = _tight(ftr.paragraphs[0], 0, 0)
    r = fp.add_run("CONFIDENTIAL \u00b7 Bander Said Allehiany for "
                   "Environmental Consultancy")
    r.font.size = Pt(7)
    r.font.color.rgb = MUT
    fp2 = _tight(ftr.add_paragraph(), 0, 0)
    fp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = fp2.add_run("Page ")
    rr.font.size = Pt(8)
    rr.font.color.rgb = NAVY
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), " PAGE ")
    fp2._p.append(fld)

    try:
        from report.template_builder import _watermark
        wm = os.path.join(ASSETS, "watermark.png")
        if os.path.exists(wm):
            _watermark(body, wm)
    except Exception:  # noqa: BLE001
        log.warning("watermark skipped", exc_info=True)

    # ---- document control ------------------------------------------------
    p = _tight(doc.add_paragraph(), 12, 2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.font.size = Pt(17)
    r.font.bold = True
    p = _tight(doc.add_paragraph(), 0, 2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("for").font.size = Pt(11)
    p = _tight(doc.add_paragraph(), 0, 8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(campaign.project_name or "")
    r.font.size = Pt(14)
    r.font.bold = True

    lab_line = settings.laboratory or "\u2014"
    if settings.lab_accreditation:
        lab_line = f"{lab_line}  \u00b7  {settings.lab_accreditation}"
    _table(doc, ["Project", campaign.project_name or "\u2014"],
           [["Document Title", title],
            ["Client", campaign.client or "\u2014"],
            ["Site", campaign.site_name or "\u2014"],
            ["Sampling date", _date_only(campaign.monitoring_start)],
            ["Reporting date",
             _date_only(getattr(campaign, "reporting_date", None))],
            ["Samples", str(len(samples))],
            ["Laboratory", lab_line],
            ["Report number", campaign.report_number or "\u2014"],
            ["Revision", campaign.revision or "00"],
            ["Prepared by", getattr(campaign, "prepared_by", None) or "\u2014"],
            ["Supervision",
             getattr(campaign, "project_supervision", None) or "\u2014"]],
           [52, 112], aligns=["left", "left"])
    _caption(doc, "Revision History")
    _table(doc, ["Revision", "Date", "Description", "Prepared", "Approved"],
           [[campaign.revision or "00",
             _date_only(getattr(campaign, "reporting_date", None)),
             "Issued for review",
             getattr(campaign, "prepared_by", None) or "\u2014", "\u2014"]],
           [22, 30, 62, 25, 25])

    # ---- contents --------------------------------------------------------
    doc.add_page_break()
    _heading(doc, "Table of Contents", 1)
    _toc_anchor(doc, r'TOC \o "1-2" \h \z \u')
    _heading(doc, "List of Tables", 1)
    _toc_anchor(doc, r'TOC \h \z \c "Table"')
    _heading(doc, "List of Figures", 1)
    _toc_anchor(doc, r'TOC \h \z \c "Figure"')

    # ---- 1. introduction -------------------------------------------------
    doc.add_page_break()
    _heading(doc, "1. Introduction", 1)
    _body(doc,
          f"Bander Said Allehiany for Environmental Consultancy was "
          f"appointed by {campaign.client or 'the client'} to undertake "
          f"{medium} quality monitoring at "
          f"{campaign.site_name or campaign.project_name or 'the site'}. "
          f"This report presents the laboratory results of that monitoring "
          f"and assesses them against the applicable national standard.")
    _body(doc,
          f"Sampling was carried out on "
          f"{_date_only(campaign.monitoring_start)}. "
          f"{len(samples)} sample(s) were recovered from the locations "
          f"listed in Section 2 and submitted for laboratory analysis. "
          f"Results were reported on "
          f"{_date_only(getattr(campaign, 'reporting_date', None))}.")
    _muted(doc,
           f"This report addresses {medium} quality only. It does not "
           f"constitute a contaminated land risk assessment, a remediation "
           f"strategy, or an opinion on geotechnical suitability.")

    # ---- 2. sampling locations -------------------------------------------
    _heading(doc, "2. Sampling locations", 1)
    _body(doc,
          "Sampling locations were selected to give coverage across the "
          "works area. Coordinates were recorded in the field at the point "
          "of sampling.")
    rows = []
    for s in samples:
        ctx = s.context
        depth = "\u2014"
        if ctx.depth_from_m is not None or ctx.depth_to_m is not None:
            depth = (f"{ctx.depth_from_m if ctx.depth_from_m is not None else 0:g}"
                     f" \u2013 "
                     f"{ctx.depth_to_m if ctx.depth_to_m is not None else 0:g} m")
        elif ctx.depth:
            depth = L.DEPTH_LABELS.get(ctx.depth, ctx.depth)
        rows.append([s.label or s.code, s.code,
                     _coord(s.latitude), _coord(s.longitude), depth,
                     s.location_name or "\u2014"])
    _table(doc, ["Sample", "Sample code", "Latitude", "Longitude", "Depth",
                 "Location"],
           rows, [18, 40, 24, 24, 26, 32],
           aligns=["center", "left", "center", "center", "center", "left"])
    _caption(doc, f"Table {tnum()} \u2014 Sampling locations and sample "
                  f"identifiers")

    if site_map_path and os.path.exists(site_map_path):
        p = _tight(doc.add_paragraph(), 4, 0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic(p, site_map_path, 150)
        _caption(doc, f"Figure {fnum()} \u2014 Sampling locations")

    if figs.get("grain_size"):
        p = _tight(doc.add_paragraph(), 4, 0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic(p, figs["grain_size"], 155)
        _caption(doc, f"Figure {fnum()} \u2014 Grain size distribution by "
                      f"sample")
        _body(doc,
              "Grain size analysis establishes the soil classification "
              "defined in Article (1) of the Executive Regulation, which "
              "determines which set of limits in Appendix (1) applies.")

    # ---- 3. methodology and basis of assessment --------------------------
    _heading(doc, "3. Methodology and basis of assessment", 1)
    _heading(doc, "3.1 Sampling and analysis", 2)
    _bullets(doc, [
        "Samples were recovered using clean equipment, decontaminated "
        "between locations.",
        "Each sample was placed in laboratory-supplied containers, "
        "labelled, and kept under chilled conditions in transit.",
        "Chain of custody documentation accompanied the samples to the "
        "laboratory.",
        "Analysis followed the methods listed against each parameter in "
        f"Appendix {_appendix['certificates']}, referencing APHA, ASTM and "
        "USEPA standards.",
    ])

    _heading(doc, "3.2 Basis of assessment", 2)
    if citation:
        _body(doc, f"Results are assessed against {citation}.")
    else:
        _body(doc,
              "No national standard has been applied to this campaign. "
              "Results are reported for information only.")

    if standard == "ncec_soil":
        _body(doc,
              "The limits in Appendix (1) are not single values: they vary "
              "by soil particle size, by sampling depth, and by the use to "
              "which the land is put. The values applied throughout this "
              "report are those selected by the context recorded below.")
    elif standard == "ncec_water_ambient":
        _body(doc,
              "The limits in Appendix (1) vary by the class of the water "
              "body sampled. The class recorded below determines the values "
              "applied throughout this report.")
    elif standard == "ncec_water_discharge":
        _body(doc,
              "The limits in Appendices (2) and (3) vary by the destination "
              "of the discharge, and are stated as averages over a period "
              "with a separate maximum for any single sample. Where a "
              "single sample is assessed, only the single-sample maximum "
              "applies.")

    ctx = settings.default_context
    ctx_rows = []
    if standard == "ncec_soil":
        ctx_rows = [
            ["Standard applied", "NCEC Appendix (1) \u2014 soil"],
            ["Soil particle size",
             L.PARTICLE_SIZE_LABELS.get(ctx.particle_size, "Not stated")],
            ["Land use", L.LAND_USE_LABELS.get(ctx.land_use, "Not stated")],
            ["Sampling depth", L.DEPTH_LABELS.get(ctx.depth, "Not stated")],
        ]
    elif standard == "ncec_water_ambient":
        ctx_rows = [
            ["Standard applied", "NCEC Appendix (1) \u2014 aqueous media"],
            ["Water body class",
             L.WATER_MEDIA_LABELS.get(ctx.water_medium, "Not stated")],
        ]
    elif standard == "ncec_water_discharge":
        ctx_rows = [
            ["Standard applied", "NCEC Appendices (2) and (3)"],
            ["Discharge destination",
             L.DISCHARGE_DESTINATION_LABELS.get(ctx.discharge_destination,
                                                "Not stated")],
            ["Assessment basis",
             "Single sample \u2014 maximum for any sample applies"
             if ctx.is_single_sample else "Averaged over the stated period"],
        ]
    if ctx_rows:
        ctx_rows.append([
            "Decision rule",
            "Simple acceptance (ILAC-G8:2019)"
            if summary.decision_rule == "simple_acceptance"
            else "Guard band \u2014 result widened by its expanded "
                 "measurement uncertainty"])
        _table(doc, ["Item", "Applied"], ctx_rows, [58, 106],
               aligns=["left", "left"])
        _caption(doc, f"Table {tnum()} \u2014 Context determining the "
                      f"applicable limits")

    _muted(doc,
           "Where the standard gives no limit for a determined parameter, "
           "the result is reported and no compliance conclusion is drawn. "
           "Results reported below the laboratory limit of quantification "
           "are treated as compliant and printed as reported.")

    # ---- 4. results ------------------------------------------------------
    doc.add_page_break()
    _heading(doc, "4. Results", 1)
    labels = [s.label for s in summary.samples]
    _body(doc,
          f"Results for {'all ' if len(labels) > 1 else ''}"
          f"{len(labels)} sample(s) are presented in Table {_tn['n'] + 1}. "
          f"Values exceeding the applicable limit are shaded.")

    headers = ["Parameter", "Unit"] + labels + ["Limit", "Result"]
    span = len(headers)
    n_s = max(1, len(labels))
    # 164 mm across: parameter 42, unit 14, limit 22, verdict 22, samples
    # share what is left, floored so a wide campaign still fits the page.
    sample_w = max(11.0, (164 - 42 - 14 - 22 - 22) / n_s)
    widths = [42, 14] + [sample_w] * n_s + [22, 22]

    rows, fills, bolds, group_at = [], [], [], []
    for entry in summary.rows:
        if entry.get("kind") == "group":
            group_at.append((len(rows), entry.get("label", "")))
            rows.append([entry.get("label", "")] + [""] * (span - 1))
            fills.append([None] * span)
            bolds.append([False] * span)
            continue
        cells = entry.get("cells") or []
        by_label = {c["sample_label"]: c for c in cells}
        verdicts = {c["verdict"] for c in cells if c.get("value") is not None
                    or c.get("below_loq")}
        if "exceeds" in verdicts:
            result_word = "Exceeds"
        elif "complies" in verdicts:
            result_word = "Complies"
        else:
            result_word = "Not assessed"
        row = [entry["analyte_name"], entry.get("unit") or "\u2014"]
        fill = [None, None]
        bold = [False, False]
        for lb in labels:
            c = by_label.get(lb)
            row.append((c or {}).get("display_value", "\u2014"))
            exceeded = bool(c and c.get("verdict") == "exceeds")
            fill.append(EXCEED_FILL if exceeded else None)
            bold.append(exceeded)
        row.append(entry.get("limit_display") or "No limit")
        fill.append(None)
        bold.append(False)
        row.append(result_word)
        fill.append(EXCEED_FILL if result_word == "Exceeds" else None)
        bold.append(result_word == "Exceeds")
        rows.append(row)
        fills.append(fill)
        bolds.append(bold)

    aligns = (["left", "center"] + ["center"] * n_s + ["center", "center"])
    table = _table(doc, headers, rows, widths, aligns=aligns,
                   fills=fills, bolds=bolds)
    for index, label in group_at:
        _group_row(table, index + 1, label, span)
    _caption(doc, f"Table {tnum()} \u2014 Analysis results by sample against "
                  f"the applicable limits")

    if figs.get("percent_of_limit"):
        p = _tight(doc.add_paragraph(), 4, 0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic(p, figs["percent_of_limit"], 155)
        _caption(doc, f"Figure {fnum()} \u2014 Results expressed as a "
                      f"percentage of the applicable limit")

    if summary.unresolved_names:
        _muted(doc,
               "The following parameters were reported by the laboratory "
               "but are not held in the parameter library, and are "
               "therefore reported without a compliance conclusion: "
               + ", ".join(summary.unresolved_names) + ".")

    # ---- 5. compliance ---------------------------------------------------
    _heading(doc, "5. Compliance assessment", 1)

    if summary.blocking_note:
        _callout(doc, summary.blocking_note, AMBER_HEX, "FFF8E7",
                 "NO VERDICT")

    v_rows, v_fills = [], []
    for e in summary.samples:
        outcome = {"compliant": "Compliant",
                   "non_compliant": "Non-compliant",
                   "no_verdict": "No verdict"}.get(e.outcome, "No verdict")
        v_rows.append([e.label, str(e.assessed_count),
                       str(e.not_assessed_count), str(e.exceedance_count),
                       outcome])
        v_fills.append([None, None, None,
                        EXCEED_FILL if e.exceedance_count else None,
                        EXCEED_FILL if e.outcome == "non_compliant" else None])
    if v_rows:
        _table(doc, ["Sample", "Parameters assessed", "Not assessed",
                     "Exceedances", "Outcome"],
               v_rows, [26, 40, 32, 30, 36], fills=v_fills)
        _caption(doc, f"Table {tnum()} \u2014 Compliance outcome by sample")

    exceedances = [c for c in summary.cells if c.verdict == "exceeds"]
    if exceedances:
        _heading(doc, "5.1 Exceedances", 2)
        items = []
        for c in exceedances:
            over = ""
            if c.limit_value and c.value is not None and c.limit_value > 0:
                pct = 100.0 * (c.value - c.limit_value) / c.limit_value
                if pct > 0:
                    over = f", an exceedance of {pct:.0f} per cent"
            items.append(
                f"{c.sample_label}: {c.analyte_name} was determined at "
                f"{c.display_value} {c.unit} against a limit of "
                f"{c.limit_display} {c.unit}{over}.")
        _bullets(doc, items)
        _callout(doc,
                 "Article (7) of the Executive Regulation requires the "
                 "National Center for Environmental Compliance to be "
                 "notified as soon as pollution is detected during an "
                 "activity.", RED_HEX, "FDEDEC", "ACTION REQUIRED")
    elif not summary.blocking_note:
        _callout(doc,
                 f"All parameters carrying a limit in the applicable "
                 f"standard comply, across all {len(summary.samples)} "
                 f"sample(s).", GREEN_HEX, "EAF5EF", "COMPLIANT")

    not_assessed = sorted({c.analyte_name for c in summary.cells
                           if c.verdict == "not_assessed"
                           and c.value is not None})
    if not_assessed:
        _heading(doc, "5.2 Parameters without an applicable limit", 2)
        _body(doc,
              "The following parameters were determined at the client's "
              "request. The applicable standard sets no limit for them and "
              "no compliance conclusion can be drawn; the results are "
              "reported for information only: " + ", ".join(not_assessed)
              + ".")

    _heading(doc, "5.3 Conclusions", 2)
    compliant = [e.label for e in summary.samples if e.outcome == "compliant"]
    failing = [e.label for e in summary.samples
               if e.outcome == "non_compliant"]
    if compliant:
        _body(doc,
              f"Sample(s) {', '.join(compliant)} comply with the applicable "
              f"limits for every parameter carrying a limit.")
    if failing:
        _body(doc,
              f"Sample(s) {', '.join(failing)} exceed the applicable limits "
              f"as set out above. Delineation sampling around the affected "
              f"location(s) is recommended to establish the extent of the "
              f"material, together with consideration of remediation "
              f"options in accordance with the Executive Regulation for the "
              f"Environmental Rehabilitation of Degraded Sites and "
              f"Remediation of Polluted Sites.")
    _muted(doc,
           "The conclusions above apply to the sampled locations and depths "
           "only. They should not be extrapolated to unsampled areas of the "
           "site.")

    if site_photo_paths:
        _heading(doc, "5.4 Site photographs", 2)
        for i, path in enumerate(site_photo_paths[:4]):
            if not (path and os.path.exists(path)):
                continue
            p = _tight(doc.add_paragraph(), 2, 0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pic(p, path, 120)
            _caption(doc, f"Figure {fnum()} \u2014 Sampling location "
                          f"{i + 1}")

    # ---- Appendix A — limits by land use ---------------------------------
    if show_land_use_table:
        doc.add_page_break()
        _heading(doc,
                 f"Appendix {_appendix['limits']} \u2014 Applicable limits, "
                 f"NCEC Appendix (1)", 1)
        _body(doc,
              "The table below reproduces the Appendix (1) limits for the "
              "parameters determined in this report, for each land use. The "
              "column applied in this assessment is shown in bold.")
        keys = [r["analyte_key"] for r in summary.rows
                if r.get("kind") == "analyte"]
        comp = land_use_comparison(keys, ctx.particle_size, ctx.depth)
        if comp:
            applied = ctx.land_use
            headers = ["Parameter", "Unit"] + [L.LAND_USE_LABELS[k]
                                               for k in L.LAND_USES]
            rows, bolds = [], []
            for r in comp:
                row = [r["analyte_name"], r["unit"]]
                bold = [False, False]
                for k in L.LAND_USES:
                    v = r["values"].get(k)
                    row.append("\u2014" if v is None else f"{v:g}")
                    bold.append(k == applied)
                rows.append(row)
                bolds.append(bold)
            _table(doc, headers, rows, [40, 14] + [22] * len(L.LAND_USES),
                   aligns=["left", "center"] + ["center"] * len(L.LAND_USES),
                   bolds=bolds)
            _caption(doc,
                     f"Table {tnum()} \u2014 NCEC Appendix (1) limits by "
                     f"land use, "
                     f"{L.PARTICLE_SIZE_LABELS.get(ctx.particle_size, '')}")
            _muted(doc,
                   "Limits vary by land use. Selecting the wrong column "
                   "produces a compliance conclusion against the wrong "
                   "standard.")

    # ---- Appendix — certificates of analysis -----------------------------
    doc.add_page_break()
    _heading(doc, f"Appendix {_appendix['certificates']} \u2014 Certificates "
                  f"of analysis", 1)
    _body(doc,
          "A certificate is reproduced for each sample, giving the method, "
          "limit of quantification and expanded measurement uncertainty for "
          "every determination.")
    for i, s in enumerate(samples):
        if not s.results:
            continue
        if i:
            doc.add_page_break()
        p = _tight(doc.add_paragraph(), 6, 2)
        r = p.add_run(f"{s.label or s.code} \u2014 {s.code}")
        r.font.size = Pt(11)
        r.font.bold = True
        r.font.color.rgb = NAVY
        meta = "  \u00b7  ".join(x for x in (
            L.PARTICLE_SIZE_LABELS.get(s.context.particle_size),
            L.DEPTH_LABELS.get(s.context.depth),
            f"Received {_date_only(s.received_at)}" if s.received_at else None,
            f"Reported {_date_only(s.reported_at)}" if s.reported_at else None,
        ) if x)
        if meta:
            _muted(doc, meta, size=9)
        rows = []
        for res in s.results:
            analyte = L.ANALYTES_BY_KEY.get(res.analyte_key)
            rows.append([
                res.reported_name or (analyte.name if analyte
                                      else res.analyte_key),
                res.raw_value or ("\u2014" if res.value is None
                                  else f"{res.value:g}"),
                res.unit or (analyte.unit if analyte else "\u2014"),
                res.method or (analyte.method if analyte else "\u2014"),
                "\u2014" if res.loq is None else f"{res.loq:g}",
                "\u2014" if res.mu_percent is None
                else f"\u00b1{res.mu_percent:g}",
            ])
        _table(doc, ["Parameter", "Result", "Unit", "Method", "LOQ", "MU %"],
               rows, [46, 22, 18, 42, 18, 18],
               aligns=["left", "center", "center", "left", "center",
                       "center"])
        _caption(doc, f"Table {tnum()} \u2014 Certificate of analysis, "
                      f"{s.label or s.code}")

    if calibration_items:
        doc.add_page_break()
        _heading(doc, "Laboratory calibration certificates", 2)
        for i, c in enumerate(calibration_items):
            path = c.get("path")
            if not (path and os.path.exists(path)):
                continue
            if i:
                doc.add_page_break()
            p = _tight(doc.add_paragraph(), 2, 4)
            r = p.add_run((c.get("title") or "Calibration certificate").strip())
            r.font.size = Pt(9.5)
            r.font.bold = True
            r.font.color.rgb = NAVY
            p = _tight(doc.add_paragraph(), 0, 0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pic(p, path, 165)

    if license_image_paths:
        doc.add_page_break()
        _heading(doc, f"Appendix {_appendix['licence']} \u2014 Environmental "
                      f"license for the institution", 1)
        for i, path in enumerate(license_image_paths):
            if not (path and os.path.exists(path)):
                continue
            if i:
                doc.add_page_break()
            p = _tight(doc.add_paragraph(), 2, 0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pic(p, path, 165)

    # ---- end marker ------------------------------------------------------
    doc.add_page_break()
    for _ in range(3):
        _tight(doc.add_paragraph(), 0, 0)
    p = _tight(doc.add_paragraph(), 0, 2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("END OF REPORT")
    r.font.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = NAVY
    p = _tight(doc.add_paragraph(), 0, 0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("  \u00b7  ".join(
        x for x in (campaign.report_number,
                    f"Rev {campaign.revision or '00'}",
                    campaign.project_name) if x))
    r.font.size = Pt(9)
    r.font.color.rgb = MUT

    os.makedirs(os.path.dirname(os.path.abspath(out_path)), exist_ok=True)
    doc.save(out_path)
    return out_path
