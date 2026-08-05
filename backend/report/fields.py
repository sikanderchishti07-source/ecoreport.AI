# -*- coding: utf-8 -*-
"""Write the report's own numbers into the document instead of asking Word.

Two jobs, both run after ``DocxTemplate.render()``:

``populate_field_caches``
    Fills each caption's SEQ field cache, so "Table 7 — Summary of SO2
    Results" reads correctly everywhere.

``build_indexes``
    Replaces the three index fields (Table of Contents, List of Figures,
    List of Tables) with real entries whose page numbers are cross-references
    to bookmarks placed on the headings and captions themselves.

Why the page numbers are cross-references
-----------------------------------------
Page numbers need a layout pass, and the only layout engine available on the
server is LibreOffice. LibreOffice and Word do not paginate identically: the
same report runs to about 63 pages in LibreOffice and about 55 in Word. So a
measured number written in as plain text is right in the PDF and wrong in the
DOCX, by a margin that grows down the document.

The number is therefore stored as a PAGEREF field pointing at a bookmark on
the heading or caption it describes, carrying the LibreOffice-measured value
as its cached result:

* Word evaluates the field against its own layout, so the DOCX is correct —
  and stays correct after a reviewer edits a paragraph, because the bookmark
  travels with the heading.
* LibreOffice resolves the cross-reference itself, or falls back to the cached
  value it measured; either way the server-produced PDF is unchanged.

``updateFields`` is switched on so Word refreshes on open. This is safe only
because ``_strip_field`` deletes the TOC field first — an earlier attempt left
a live TOC field in place and Word regenerated a contents page reading 1
against every entry. There is now no TOC field to regenerate, only
cross-references to fixed bookmarks.

Caption numbers stay as computed SEQ caches: they are already verified correct
and Word recomputes the same sequence.

The index entries are inserted *before* the measuring conversion with their
numbers left blank, so the index pages are already at full height when
pagination is measured; filling in a two-digit number afterwards cannot push
anything onto a different page.
"""
from __future__ import annotations

import itertools
import logging
import os
import re
import shutil
import subprocess
import tempfile
import zipfile
from typing import Dict, List, Optional, Tuple

log = logging.getLogger(__name__)

DOC_XML = "word/document.xml"
SETTINGS_XML = "word/settings.xml"

XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"

# One field = one run: begin, instrText, separate, cached result, end.
_FIELD_RE = re.compile(
    r'(<w:fldChar\s+w:fldCharType="begin"[^/]*/>'
    r'<w:instrText[^>]*>)(.*?)(</w:instrText>'
    r'<w:fldChar\s+w:fldCharType="separate"\s*/>)'
    r'<w:t[^>]*>[^<]*</w:t>'
    r'(<w:fldChar\s+w:fldCharType="end"\s*/>)',
    re.S,
)
_SEQ_RE = re.compile(r"\bSEQ\s+(\w+)", re.I)

INDEX_TITLES = ("Table of Contents", "List of Figures", "List of Tables")

# Elements that must follow w:updateFields inside w:settings. The tag is
# inserted before whichever of these appears first.
SETTINGS_AFTER_UPDATE = (
    "w:hdrShapeDefaults", "w:footnotePr", "w:endnotePr", "w:compat",
    "w:docVars", "w:rsids", "m:mathPr", "w:themeFontLang",
    "w:clrSchemeMapping", "w:doNotAutoCompressPictures", "w:shapeDefaults",
    "w:decimalSymbol", "w:listSeparator",
)

# Word reserves ids below 1000 for its own bookmarks in some documents; start
# clear of them and never reuse a number within one build.
_BOOKMARK_IDS = itertools.count(9000)


# ---------------------------------------------------------------------------
# Shared
# ---------------------------------------------------------------------------
def _rewrite_zip(path: str, names: List[str], blobs: Dict[str, bytes]) -> None:
    fd, tmp = tempfile.mkstemp(suffix=".docx",
                               dir=os.path.dirname(os.path.abspath(path)))
    os.close(fd)
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for n in names:                    # preserve original entry order
            zout.writestr(n, blobs[n])
    shutil.move(tmp, path)


def _norm(text: str) -> str:
    """Collapse whitespace and drop soft hyphens so PDF text and document text
    compare equal even where the layout engine has rewrapped a line."""
    return re.sub(r"\s+", " ", (text or "").replace("\u00ad", "")).strip()


# ---------------------------------------------------------------------------
# 1. Caption numbers
# ---------------------------------------------------------------------------
def _populate(xml: str) -> Tuple[str, Dict[str, int]]:
    counters: Dict[str, int] = {}
    out: List[str] = []
    last = 0
    for m in _FIELD_RE.finditer(xml):
        instr = m.group(2)
        seq = _SEQ_RE.search(instr)
        if not seq:
            continue                       # TOC / PAGE: handled elsewhere
        kind = seq.group(1)
        counters[kind] = counters.get(kind, 0) + 1
        out.append(xml[last:m.start()])
        out.append(m.group(1) + instr + m.group(3)
                   + "<w:t>%d</w:t>" % counters[kind] + m.group(4))
        last = m.end()
    out.append(xml[last:])
    return "".join(out), counters


def populate_field_caches(docx_path: str) -> Dict[str, int]:
    """Write computed caption numbers into their field caches, in place."""
    try:
        with zipfile.ZipFile(docx_path) as zin:
            names = zin.namelist()
            if DOC_XML not in names:
                return {}
            blobs = {n: zin.read(n) for n in names}

        fixed, counters = _populate(blobs[DOC_XML].decode("utf-8"))
        if not counters:
            return {}
        blobs[DOC_XML] = fixed.encode("utf-8")
        _rewrite_zip(docx_path, names, blobs)
        log.info("caption numbers written: %s",
                 ", ".join(f"{k} 1-{v}" for k, v in sorted(counters.items())))
        return counters
    except Exception:  # noqa: BLE001
        log.exception("caption numbering failed — captions may be blank")
        return {}


# ---------------------------------------------------------------------------
# 2. Bookmarks and cross-reference fields
# ---------------------------------------------------------------------------
def _add_bookmark(paragraph, name: str) -> None:
    """Wrap a paragraph in a hidden bookmark.

    The marker is zero-width and carries no formatting, so it cannot affect
    where anything falls on the page. Names beginning with an underscore are
    hidden from Word's bookmark list, which is what Word uses for its own
    cross-reference targets.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    bid = str(next(_BOOKMARK_IDS))
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), bid)
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bid)

    p = paragraph._p
    pPr = p.find(qn("w:pPr"))
    if pPr is not None:
        pPr.addnext(start)                 # must follow the properties block
    else:
        p.insert(0, start)
    p.append(end)


def _pageref_field(paragraph, bookmark: str, half_pt: str = "20"):
    """Append a PAGEREF cross-reference and return its cached-result element.

    Five runs, as Word writes them: begin, instruction, separate, cached
    result, end. Word replaces the cached result with the real page number;
    readers that do not evaluate fields display the cached value as-is.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    def _run():
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        for tag in ("w:sz", "w:szCs"):
            el = OxmlElement(tag)
            el.set(qn("w:val"), half_pt)
            rPr.append(el)
        r.append(rPr)
        return r

    begin = _run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    # A dirty field is refreshed by Word when the document is opened. This is
    # independent of the settings file, so the numbers come right even if
    # updateFields is ignored. LibreOffice ignores w:dirty, so the PDF is
    # unaffected.
    fld.set(qn("w:dirty"), "true")
    begin.append(fld)

    instr_run = _run()
    instr = OxmlElement("w:instrText")
    instr.set(XML_SPACE, "preserve")
    instr.text = " PAGEREF %s \\h " % bookmark
    instr_run.append(instr)

    sep = _run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "separate")
    sep.append(fld)

    cache_run = _run()
    cache = OxmlElement("w:t")
    cache.text = ""                        # filled in after measuring
    cache_run.append(cache)

    end = _run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "end")
    end.append(fld)

    for r in (begin, instr_run, sep, cache_run, end):
        paragraph._p.append(r)
    return cache


# ---------------------------------------------------------------------------
# 3. Index pages
# ---------------------------------------------------------------------------
def _pdf_pages(pdf_path: str) -> List[str]:
    """Per-page plain text. PyMuPDF is already a dependency; pdftotext is the
    fallback for hosts where it cannot be imported."""
    try:
        import fitz                        # PyMuPDF
        with fitz.open(pdf_path) as doc:
            return [p.get_text() for p in doc]
    except Exception:  # noqa: BLE001
        log.info("PyMuPDF unavailable — falling back to pdftotext")
    try:
        out = subprocess.run(["pdftotext", "-layout", pdf_path, "-"],
                             capture_output=True, timeout=180, check=True)
        return out.stdout.decode("utf-8", "replace").split("\f")
    except Exception:  # noqa: BLE001
        log.exception("could not read page text from %s", pdf_path)
        return []


def _page_lookup(pages: List[str]) -> Tuple[List[str], set]:
    """Normalised page text, plus the indices of the index pages themselves.

    Those pages repeat every heading and caption verbatim, so they must be
    excluded or every entry would resolve to the contents page. Detection is
    by dot leaders rather than by title, so it works just as well on the
    Arabic half of a bilingual report.
    """
    norm = [_norm(p) for p in pages]
    skip = set()
    for i, (raw, t) in enumerate(zip(pages, norm)):
        if any(title in t for title in INDEX_TITLES):
            skip.add(i)
        elif raw.count("....") >= 5:          # a page of dot-leader entries
            skip.add(i)
    return norm, skip


def _find_page(norm_pages: List[str], skip: set, text: str) -> Optional[int]:
    needle = _norm(text)
    if not needle:
        return None
    for i, page in enumerate(norm_pages):
        if i not in skip and needle in page:
            return i + 1
    # Long captions sometimes wrap awkwardly; retry on a distinctive prefix.
    short = " ".join(needle.split()[:6])
    if len(short) >= 12:
        for i, page in enumerate(norm_pages):
            if i not in skip and short in page:
                return i + 1
    return None


def _plan(doc) -> List[Tuple[object, List[Tuple[int, str, str]]]]:
    """Pair every index anchor with the entries that belong to it.

    A bilingual report contains the whole English document followed by the
    whole Arabic one, so there are two Tables of Contents, two Lists of
    Figures and two Lists of Tables. Walking the document once and assigning
    each heading or caption to the most recent anchor of its kind keeps each
    index describing its own half, instead of the English index absorbing the
    Arabic captions as well.

    Each heading and caption also gets its bookmark here, so the index entry
    built later has something to point at.
    """
    buckets: Dict[int, List[Tuple[int, str, str]]] = {}
    order: List[object] = []
    current: Dict[str, object] = {"toc": None, "fig": None, "tab": None}
    marks = itertools.count(1)

    def _open(kind: str, paragraph) -> None:
        current[kind] = paragraph
        buckets[id(paragraph)] = []
        order.append(paragraph)

    def _mark(paragraph) -> str:
        name = "_EcoRef%04d" % next(marks)
        _add_bookmark(paragraph, name)
        return name

    for p in doc.paragraphs:
        xml = p._p.xml
        if "instrText" in xml and "TOC" in xml:
            if "\\o" in xml:
                _open("toc", p)
            elif "Figure" in xml:
                _open("fig", p)
            elif "Table" in xml:
                _open("tab", p)
            continue

        text = (p.text or "").strip()
        if not text:
            continue
        style = (p.style.name or "") if p.style is not None else ""

        if style.startswith("Heading"):
            if text in INDEX_TITLES:
                continue
            anchor = current["toc"]
            if anchor is None:
                continue
            m = re.search(r"(\d)", style)
            level = int(m.group(1)) if m else 1
            if level <= 3:
                buckets[id(anchor)].append((level, text, _mark(p)))
        elif style == "Caption":
            kind = "fig" if text.startswith("Figure") else (
                "tab" if text.startswith("Table") else None)
            if kind is None:
                continue
            anchor = current[kind]
            if anchor is not None:
                buckets[id(anchor)].append((1, text, _mark(p)))

    return [(a, buckets[id(a)]) for a in order if buckets[id(a)]]


def _entry(doc, level: int, text: str, bookmark: str, right_tab_in: float):
    """One index line: text, dot leader, right-aligned page-number field."""
    from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
    from docx.shared import Inches, Pt

    p = doc.add_paragraph()                # appended, then moved into place
    pf = p.paragraph_format
    pf.left_indent = Inches(0.24 * (level - 1))
    # Index entries inherit Normal, which carries body-text leading and six
    # points after each paragraph. Over thirty-four contents entries that
    # spacing alone costs about four centimetres — enough to push the last
    # few onto a second, nearly empty page. Single leading and a hairline gap
    # keep the list readable and fit it on one sheet.
    pf.space_before = Pt(0)
    pf.space_after = Pt(1.5)
    pf.line_spacing = 1.0
    pf.tab_stops.add_tab_stop(Inches(right_tab_in), WD_TAB_ALIGNMENT.RIGHT,
                              WD_TAB_LEADER.DOTS)
    run = p.add_run(text)
    run.font.size = Pt(10)
    tab = p.add_run("\t")                  # dot leader to the right margin
    tab.font.size = Pt(10)
    cache = _pageref_field(p, bookmark)    # number resolved by the reader
    return p, cache


def _insert_after(anchor, paragraphs) -> None:
    """Place paragraphs directly after the anchor, keeping their order."""
    for p in reversed(paragraphs):
        anchor._p.addnext(p._p)


def _strip_field(paragraph) -> None:
    """Remove the now-redundant TOC field so Word cannot regenerate a second,
    duplicate index on top of the one we just wrote. This also makes it safe
    to ask Word to refresh fields on open."""
    for run in list(paragraph.runs):
        if "fldChar" in run._r.xml or "instrText" in run._r.xml:
            run._r.getparent().remove(run._r)


def _set_update_on_open(docx_path: str, enabled: bool = True) -> None:
    """Ask Word to refresh fields when the document is opened.

    Every page number is now a cross-reference to a bookmark, so a refresh is
    exactly what makes the contents agree with the pages in Word — including
    after a reviewer has edited the text. The TOC fields have been removed by
    this point, so there is nothing left that a refresh could regenerate
    incorrectly.

    LibreOffice ignores this setting, so the PDF path is unaffected.
    """
    try:
        with zipfile.ZipFile(docx_path) as zin:
            names = zin.namelist()
            blobs = {n: zin.read(n) for n in names}
        if SETTINGS_XML not in blobs:
            return
        settings = blobs[SETTINGS_XML].decode("utf-8")
        cleaned = re.sub(r"<w:updateFields[^/>]*/>", "", settings)
        if enabled:
            tag = '<w:updateFields w:val="true"/>'
            # settings.xml is a fixed sequence, not a bag of elements. Word
            # silently discards anything out of order — appending the tag at
            # the end looks right and does nothing at all. It belongs ahead of
            # rsids, compat, mathPr and everything after them.
            pos = None
            for name in SETTINGS_AFTER_UPDATE:
                m = re.search(r"<%s[ />]" % re.escape(name), cleaned)
                if m and (pos is None or m.start() < pos):
                    pos = m.start()
            if pos is not None:
                cleaned = cleaned[:pos] + tag + cleaned[pos:]
            elif "</w:settings>" in cleaned:
                cleaned = cleaned.replace("</w:settings>", tag + "</w:settings>")
            else:                          # self-closing, empty settings part
                cleaned = re.sub(r"(<w:settings[^>]*)/>",
                                 r"\1>" + tag + "</w:settings>", cleaned)
        if cleaned != settings:
            blobs[SETTINGS_XML] = cleaned.encode("utf-8")
            _rewrite_zip(docx_path, names, blobs)
    except Exception:  # noqa: BLE001
        log.exception("could not set updateFields")


def build_indexes(docx_path: str, convert_to_pdf) -> bool:
    """Populate the three index pages with real entries and page numbers.

    ``convert_to_pdf(docx_path, out_dir) -> pdf_path`` is passed in rather
    than imported, so this module stays independent of the report package's
    conversion strategy.

    Returns True when page numbers were resolved. On any failure the document
    still keeps its entries — an index without page numbers is a great deal
    better than three blank pages, and Word will fill the numbers in on open
    from the bookmarks regardless.
    """
    try:
        from docx import Document

        doc = Document(docx_path)
        plan = _plan(doc)
        if not plan:
            log.warning("no headings or captions found — indexes left empty")
            return False

        # Length arithmetic returns a plain int (EMU) in python-docx, so the
        # conversion to inches is done explicitly. The first section can carry
        # zero margins (the cover is laid out full-bleed), which would push the
        # tab stop past the printable area and wrap every page number onto its
        # own line — so use the widest sensible margin defined anywhere.
        EMU_IN = 914400.0
        usable = []
        for sec in doc.sections:
            width = int(sec.page_width or 0)
            left = int(sec.left_margin or 0)
            right = int(sec.right_margin or 0)
            if width and left and right:
                usable.append((width - left - right) / EMU_IN)
        if usable:
            right_tab = min(usable)
        else:
            page_in = int(doc.sections[0].page_width or 0) / EMU_IN or 8.27
            right_tab = page_in - 2.0       # assume 1" margins
        right_tab = max(right_tab, 2.0)

        caches: List[Tuple[object, str]] = []
        for anchor, entries in plan:
            built = []
            for level, text, bookmark in entries:
                p, cache = _entry(doc, level, text, bookmark, right_tab)
                built.append(p)
                caches.append((cache, text))
            _insert_after(anchor, built)
            _strip_field(anchor)

        doc.save(docx_path)                # entries present, numbers blank

        # Measure. The index pages are already at full height, so the page a
        # heading falls on cannot move when its number is filled in.
        with tempfile.TemporaryDirectory(prefix="ecoreport_idx_") as td:
            probe = os.path.join(td, "probe.docx")
            shutil.copy(docx_path, probe)
            try:
                pdf = convert_to_pdf(probe, td)
            except Exception:  # noqa: BLE001
                log.exception("index measuring pass failed — "
                              "entries kept without cached page numbers")
                _set_update_on_open(docx_path)
                return False
            norm_pages, skip = _page_lookup(_pdf_pages(pdf))

        if not norm_pages:
            _set_update_on_open(docx_path)
            return False

        resolved = 0
        for cache, text in caches:
            page = _find_page(norm_pages, skip, text)
            if page:
                cache.text = str(page)
                resolved += 1

        doc.save(docx_path)
        _set_update_on_open(docx_path)
        log.info("indexes built: %d index block(s), %d entries, "
                 "%d/%d page numbers cached (Word resolves from bookmarks)",
                 len(plan), len(caches), resolved, len(caches))
        return resolved > 0
    except Exception:  # noqa: BLE001
        log.exception("index build failed — index pages may be empty")
        return False
