"""Builds the master docxtpl template (master_template.docx) programmatically.

The template replicates the BSA gold-standard AAQ report structure:
cover, document control, TOC/LoF/LoT, definitions, executive summary,
sections 1-6, appendices — with Jinja placeholders for everything dynamic.

Rebuild any time with:  python -m report.template_builder
"""
from __future__ import annotations

import os

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph

ASSETS = os.path.join(os.path.dirname(__file__), "assets")

# Accreditation marks printed across the cover masthead, left to right.
# Each entry is a PNG filename in assets/ plus the caption beneath it.
#
# A badge whose file is not present is skipped entirely, so the row shows
# only what BSA actually holds and never leaves a placeholder claiming an
# accreditation the lab does not have. To change a mark, replace its PNG;
# to change the wording or a certificate number, edit it here.
ACCREDITATION_BADGES = [
    ("badge_1.png", ["ONES INTERNATIONAL LIMITED",
                     "Health and Safety Management System",
                     "Certificate No: C00175-01"]),
    ("badge_2.png", ["ACCREDITED",
                     "Management System Certification Body",
                     "MSCB-127"]),
    ("badge_3.png", ["ACCREDITED",
                     "Management System Certification Body",
                     "MSCB-127"]),
    ("badge_4.png", []),
]
OUT = os.path.join(os.path.dirname(__file__), "master_template.docx")

# Brand palette — blue-dominant with a restrained green accent
NAVY = RGBColor(0x0F, 0x3D, 0x6E)
BLUE = RGBColor(0x1F, 0x6F, 0xB2)
GREEN = RGBColor(0x2F, 0x9E, 0x63)
DARK = RGBColor(0x1F, 0x1F, 0x1F)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
NAVY_FILL = "0F3D6E"
BLUE_FILL = "1F6FB2"
SKY_FILL = "E8F1F9"
GREEN_ACCENT = RGBColor(0x2F, 0x7D, 0x32)
ZEBRA_FILL = "F4F7FA"
HERO_EDGE_FILL = "163259"    # matches the scrim at the hero's left edge        # very light blue-grey for alternating table rows
OK_GREEN = RGBColor(0x1E, 0x7D, 0x4F)     # compliant
WARN_AMBER = RGBColor(0xB0, 0x6A, 0x00)   # information only
BAD_RED = RGBColor(0xB3, 0x1F, 0x1F)      # exceedance
CALLOUT_FILL = "EEF4FA"      # key-finding panel
CALLOUT_EDGE = "1F6FB2"
RULE_GREY = "D9E1E8"
MUTED_GREY = RGBColor(0x6B, 0x6B, 0x6B)
GRAY_FILL = NAVY_FILL   # legacy alias: all header cells now use the navy fill
_DARK_FILLS = {NAVY_FILL, BLUE_FILL}


# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------
def _field(paragraph, instr: str, cached: str = " "):
    """Insert a Word field (TOC, PAGE, SEQ...) into a paragraph.

    `cached` is the field's stored result — what any reader that does not
    evaluate fields will display. LibreOffice (our PDF path) never evaluates
    them, so report.fields rewrites these caches after rendering. Leaving the
    cache blank is what produced "Table . \u2014" in every exported PDF.
    """
    r = paragraph.add_run()
    for el, attrs, text in (
        ("w:fldChar", {"w:fldCharType": "begin"}, None),
        ("w:instrText", {"xml:space": "preserve"}, instr),
        ("w:fldChar", {"w:fldCharType": "separate"}, None),
        ("w:t", {}, cached),
        ("w:fldChar", {"w:fldCharType": "end"}, None),
    ):
        e = OxmlElement(el)
        for k, v in attrs.items():
            e.set(qn(k), v)
        if text is not None:
            e.text = text
        r._r.append(e)


def _update_fields_on_open(doc):
    settings = doc.settings.element
    upd = OxmlElement("w:updateFields")
    upd.set(qn("w:val"), "true")
    settings.append(upd)


def _table_borders(table, colour=RULE_GREY, size="4"):
    """Replace the heavy default grid with hairline rules."""
    tblPr = table._tbl.tblPr
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    b = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), size)
        e.set(qn("w:color"), colour)
        b.append(e)
    tblPr.append(b)


def _zebra(table, start_row=1, fill=ZEBRA_FILL):
    """Separate alternate body rows so wide tables stay readable.

    Originally this tinted every second row. Word shading is opaque, so those
    150-odd shaded cells punched holes in the watermark and the mark appeared
    to break up wherever a table sat. Alternate rows are now separated by a
    hairline instead: the same job of guiding the eye across a wide row, done
    without blocking anything behind it.
    """
    for i, row in enumerate(table.rows[start_row:], start=start_row):
        if (i - start_row) % 2 == 1:
            for c in row.cells:
                tcPr = c._tc.get_or_add_tcPr()
                borders = tcPr.find(qn("w:tcBorders"))
                if borders is None:
                    borders = OxmlElement("w:tcBorders")
                    tcPr.append(borders)
                for side in ("top", "bottom"):
                    e = OxmlElement("w:%s" % side)
                    e.set(qn("w:val"), "single")
                    e.set(qn("w:sz"), "4")
                    e.set(qn("w:color"), "E3EAF1")
                    borders.append(e)


def _cell_pad(table, top=0.10, bottom=0.10, left=0.16, right=0.16):
    """Breathing room inside every cell — the biggest single readability win."""
    tblPr = table._tbl.tblPr
    mar = OxmlElement("w:tblCellMar")
    for tag, val in (("top", top), ("bottom", bottom),
                     ("left", left), ("right", right)):
        e = OxmlElement(f"w:{tag}")
        e.set(qn("w:w"), str(int(val * 567)))
        e.set(qn("w:type"), "dxa")
        mar.append(e)
    tblPr.append(mar)


def _polish(table, header_rows=1, zebra=True):
    """Apply the house table style: hairlines, padding, zebra striping."""
    _table_borders(table)
    _cell_pad(table)
    if zebra:
        _zebra(table, start_row=header_rows)
    return table


def _callout(doc, title, body_key, fill=CALLOUT_FILL, edge=CALLOUT_EDGE):
    """Tinted panel with a coloured left edge, for a key statement."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.rows[0].cells[0]
    tcPr = c._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)
    bdrs = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:color"), edge)
    bdrs.append(left)
    for edge_name in ("top", "bottom", "right"):
        e = OxmlElement(f"w:{edge_name}")
        e.set(qn("w:val"), "nil")
        bdrs.append(e)
    tcPr.append(bdrs)
    mar = OxmlElement("w:tcMar")
    for tag, val in (("top", 0.22), ("bottom", 0.22), ("left", 0.35),
                     ("right", 0.3)):
        e = OxmlElement(f"w:{tag}")
        e.set(qn("w:w"), str(int(val * 567)))
        e.set(qn("w:type"), "dxa")
        mar.append(e)
    tcPr.append(mar)
    p0 = c.paragraphs[0]
    p0.paragraph_format.space_after = Pt(3)
    r = p0.add_run(title)
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = NAVY
    p1 = c.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p1.paragraph_format.space_after = Pt(0)
    r1 = p1.add_run(body_key)
    r1.font.size = Pt(10)
    return t


def _shade(cell, fill=GRAY_FILL):
    if fill in _DARK_FILLS:
        for p in cell.paragraphs:
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            for r in p.runs:
                r.font.color.rgb = WHITE
                r.bold = True
                # slight letter-spacing gives the masthead row a designed feel
                rPr = r._r.get_or_add_rPr()
                sp = OxmlElement("w:spacing")
                sp.set(qn("w:val"), "8")
                rPr.append(sp)
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def _cell_text(cell, text, bold=False, size=10, align="left", italic=False):
    """Write a cell's text. Header cells (bold + shaded) are set slightly
    smaller and letter-spaced by the caller for a cleaner masthead row."""
    cell.paragraphs[0].text = ""
    p = cell.paragraphs[0]
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT,
                   "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return cell


def _merge_row(table, row_idx, start, end):
    a = table.cell(row_idx, start)
    b = table.cell(row_idx, end)
    a.merge(b)
    return a


def _p(doc, text="", size=11, bold=False, italic=False, align="left",
       style=None, space_after=6, color=None):
    p = doc.add_paragraph(style=style)
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT,
                   "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT,
                   "justify": WD_ALIGN_PARAGRAPH.JUSTIFY}[align]
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        if color:
            r.font.color.rgb = color
    return p


def _heading(doc, text, level=1):
    # note: keep_with_next below binds the heading to the paragraph that
    # follows; keep_together stops a two-line heading breaking across pages
    """Section heading. Level 1 gets a navy number badge and a full-width rule
    so sections are visually separated; deeper levels are quieter."""
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.space_before = Pt(16 if level == 1 else 11)
    p.paragraph_format.space_after = Pt(6 if level == 1 else 4)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.keep_together = True
    if level == 1:
        num, _, rest = text.partition(" ")
        if num.rstrip(".").replace(".", "").isdigit():
            r0 = p.add_run(num.rstrip(".") + "  ")
            r0.font.color.rgb = BLUE
            r0.bold = True
            r1 = p.add_run(rest)
            r1.bold = True
        else:
            p.add_run(text).bold = True
        pPr = p._p.get_or_add_pPr()
        bdr = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single")
        bot.set(qn("w:sz"), "10")
        bot.set(qn("w:space"), "4")
        bot.set(qn("w:color"), BLUE_FILL)
        bdr.append(bot)
        pPr.append(bdr)
    else:
        p.add_run(text).bold = True
    return p


def _caption(doc, kind: str, text: str):
    """'Table 5.2 — Summary of SO2 Results'.

    A single SEQ field numbers captions consecutively through the document,
    as the gold-standard report does. Word renumbers automatically; the cached
    result is rewritten by report.fields after rendering so that the PDF —
    which is produced by LibreOffice and never evaluates fields — matches.
    """
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.keep_with_next = True
    r = p.add_run(f"{kind} ")
    r.bold = True
    # Flat numbering (Table 1..15, Figure 1..20), matching the gold standard.
    # The previous "STYLEREF 1 \\s" prefix asked Word for the list number of
    # the nearest Heading 1 -- but the headings carry typed numbers, not list
    # numbering, so Word had no number to return and the prefix rendered as an
    # error or a blank.
    _field(p, f" SEQ {kind} \\* ARABIC ")
    for run in p.runs:
        run.font.color.rgb = NAVY
        run.font.size = Pt(10)
        run.bold = True
        run.italic = False
    r2 = p.add_run(f" \u2014 {text}")
    r2.bold = True
    r2.font.size = Pt(10)
    r2.font.color.rgb = DARK
    return p


def _figure(doc, key: str, caption_text: str, explain: bool = True):
    """A figure and its caption, both conditional on the figure existing.

    A chart is only produced when the parameter has valid readings, so the
    context value is empty when it has none. Printing the image placeholder
    and its caption regardless gave the reader a caption above blank space,
    consumed a figure number, and listed a figure that is not in the report.

    ``explain`` prints one line in place of the figure. Silence would leave a
    client wondering whether a page had been lost in production; the line
    says the omission was deliberate and agrees with the N/R entry the
    summary table already carries for the same parameter.
    """
    _p(doc, "{%%p if %s %%}" % key, size=1, space_after=0)
    _p(doc, "{{ %s }}" % key, align="center")
    _caption(doc, "Figure", caption_text)
    if explain:
        _p(doc, "{%p else %}", size=1, space_after=0)
        _p(doc, "No chart is presented for this parameter: no valid readings "
                "were recorded during the survey period.", size=9,
           italic=True, color=MUTED_GREY)
    _p(doc, "{%p endif %}", size=1, space_after=0)


def _summary_table(doc, rows, ncec_cols):
    """Gold-standard pollutant summary table.
    rows: list of (label, value_placeholder) tuples.
    ncec_cols: list of (period_label, limit_placeholder)."""
    n_ncec = max(len(ncec_cols), 1)
    tbl = doc.add_table(rows=2 + len(rows), cols=2 + n_ncec)
    tbl.style = "Table Grid"
    _polish(tbl)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row 0: descriptor merged | "NCEC Exceedance Level µg/m³" merged
    _merge_row(tbl, 0, 0, 1)
    _cell_text(tbl.cell(0, 0), "", size=10)
    hdr = tbl.cell(0, 2)
    if n_ncec > 1:
        hdr = _merge_row(tbl, 0, 2, 1 + n_ncec)
    _cell_text(hdr, "NCEC Exceedance Level µg/m³", bold=True, size=10, align="center")
    _shade(hdr)
    # Header row 1: blank | blank | period labels
    _merge_row(tbl, 1, 0, 1)
    for j, (period, _) in enumerate(ncec_cols):
        _cell_text(tbl.cell(1, 2 + j), period, bold=True, size=10, align="center")
        _shade(tbl.cell(1, 2 + j))
    # Data rows; NCEC limit cells merged vertically across all data rows
    for i, (label, value) in enumerate(rows):
        _cell_text(tbl.cell(2 + i, 0), label, size=10)
        _cell_text(tbl.cell(2 + i, 1), value, size=10, align="center")
    for j, (_, limit_ph) in enumerate(ncec_cols):
        top = tbl.cell(2, 2 + j)
        bottom = tbl.cell(1 + len(rows), 2 + j)
        merged = top.merge(bottom)
        _cell_text(merged, limit_ph, size=10, align="center")
    return tbl


def _tr_tag_row(table, row_idx, tag):
    """A dedicated row containing only a {%tr %} tag (docxtpl removes the row)."""
    cell = table.cell(row_idx, 0)
    cell.paragraphs[0].text = tag


def _header_footer(section):
    """Compact running header and footer.

    The previous header stacked a 1.6 cm logo, a three-line centred title and
    a second emblem, which cost close to three centimetres off the top of
    every page. Here the mark is small on the left and the document identifies
    itself on the right in two lines — title, then project, report number and
    revision — over a hairline rule. Same information, roughly a third of the
    height, and the same shape a consultancy report normally carries.
    """
    hdr = section.header
    hdr.is_linked_to_previous = False
    tbl = hdr.add_table(rows=1, cols=2, width=Cm(17))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    left, right = tbl.rows[0].cells
    left.width, right.width = Cm(4.2), Cm(12.8)

    lp = left.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lp.paragraph_format.space_after = Pt(0)
    try:
        lp.add_run().add_picture(os.path.join(ASSETS, "logo_left.png"),
                                 height=Cm(0.85))
    except Exception:
        r = lp.add_run("BSA.lab")
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = NAVY

    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_after = Pt(0)
    r1 = rp.add_run("Ambient Air Quality Monitoring Report")
    r1.bold = True
    r1.font.size = Pt(9.5)
    r1.font.color.rgb = NAVY

    rp2 = right.add_paragraph()
    rp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp2.paragraph_format.space_after = Pt(0)
    # project, report number and revision on one line: enough for a reader to
    # identify a loose page without a second trip to the cover
    r2 = rp2.add_run("{{ project_name }}  \u00b7  {{ report_number }}"
                     "  \u00b7  Rev {{ revision }}")
    r2.font.size = Pt(7.5)
    r2.font.color.rgb = MUTED_GREY

    rule = hdr.add_paragraph()
    rule.paragraph_format.space_before = Pt(1)
    rule.paragraph_format.space_after = Pt(0)
    for run in rule.runs:
        run.font.size = Pt(1)
    pPr = rule._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "8")
    bot.set(qn("w:color"), BLUE_FILL)
    borders.append(bot)
    pPr.append(borders)

    # Faded mark behind the body text. Uses assets/watermark.png when one is
    # supplied, otherwise the BSA logo faded down; absent either, skipped.
    wm_src = os.path.join(ASSETS, "watermark.png")
    if not os.path.exists(wm_src):
        logo = os.path.join(ASSETS, "logo_left.png")
        if os.path.exists(logo):
            try:
                wm_src = _make_watermark_png(
                    logo, os.path.join(ASSETS, "_watermark_generated.png"))
            except Exception:  # noqa: BLE001
                wm_src = ""
        else:
            wm_src = ""
    if wm_src and os.path.exists(wm_src):
        _watermark(section, wm_src)

    ftr = section.footer
    ftr.is_linked_to_previous = False
    ft = ftr.add_table(rows=1, cols=2, width=Cm(17))
    ft.alignment = WD_TABLE_ALIGNMENT.CENTER
    fl, fr = ft.rows[0].cells
    fl.width, fr.width = Cm(12.8), Cm(4.2)

    flp = fl.paragraphs[0]
    flp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    flp.paragraph_format.space_after = Pt(0)
    fpr = flp._p.get_or_add_pPr()
    fBdr = OxmlElement("w:pBdr")
    ftop = OxmlElement("w:top")
    ftop.set(qn("w:val"), "single")
    ftop.set(qn("w:sz"), "6")
    ftop.set(qn("w:color"), "C9D6E2")
    fBdr.append(ftop)
    fpr.append(fBdr)
    fl_run = flp.add_run("CONFIDENTIAL  \u00b7  {{ provider_legal_name }}")
    fl_run.font.size = Pt(7)
    fl_run.font.color.rgb = MUTED_GREY

    frp = fr.paragraphs[0]
    frp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    frp.paragraph_format.space_after = Pt(0)
    fpr2 = frp._p.get_or_add_pPr()
    fBdr2 = OxmlElement("w:pBdr")
    ftop2 = OxmlElement("w:top")
    ftop2.set(qn("w:val"), "single")
    ftop2.set(qn("w:sz"), "6")
    ftop2.set(qn("w:color"), "C9D6E2")
    fBdr2.append(ftop2)
    fpr2.append(fBdr2)
    pre = frp.add_run("Page ")
    pre.font.size = Pt(8)
    pre.font.color.rgb = NAVY
    _field(frp, " PAGE ")
    for r in frp.runs:
        r.font.size = Pt(8)
        r.font.color.rgb = NAVY


# ---------------------------------------------------------------------------
# Document assembly
# ---------------------------------------------------------------------------

def _modernise_settings(doc):
    """Drop the legacy compatibility flag and force field update on open, so
    Word populates the Table of Contents / List of Figures / List of Tables
    without the reader pressing Ctrl+A F9."""
    settings = doc.settings.element
    for tag in ("w:compat",):
        el = settings.find(qn(tag))
        if el is not None:
            settings.remove(el)
    if settings.find(qn("w:updateFields")) is None:
        uf = OxmlElement("w:updateFields")
        uf.set(qn("w:val"), "true")
        settings.append(uf)



def _verdict_box(doc, body_key):
    """Single-line finding under a pollutant table, in a tinted panel with a
    green left edge and a tick — the reader's takeaway for that section."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.rows[0].cells[0]
    tcPr = c._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F1F7F2")
    tcPr.append(shd)
    bdrs = OxmlElement("w:tcBorders")
    for side, colour, sz in (("left", "2F7D32", "18"), ("top", "DDE7DE", "4"),
                             ("bottom", "DDE7DE", "4"), ("right", "DDE7DE", "4")):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:color"), colour)
        bdrs.append(el)
    tcPr.append(bdrs)
    _cell_pad(t, top=0.14, bottom=0.14, left=0.28, right=0.22)
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    tick = p.add_run("\u2713  ")
    tick.bold = True
    tick.font.size = Pt(10)
    tick.font.color.rgb = OK_GREEN
    r = p.add_run(body_key)
    r.font.size = Pt(9.5)
    r.font.color.rgb = DARK
    _p(doc, space_after=6)
    return t




def _no_word_split(table):
    """Stop Word breaking a word across lines inside table cells.

    Word will hyphenate-by-force inside a narrow cell, producing headers like
    "APPLICABL / E LIMIT". Turning off automatic hyphenation for the runs in
    the table makes it wrap at spaces instead, which is what a typesetter
    would do.
    """
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                pPr = p._p.get_or_add_pPr()
                if pPr.find(qn("w:suppressAutoHyphens")) is None:
                    pPr.append(OxmlElement("w:suppressAutoHyphens"))



def _hairline(cell, colour="DCE5EE"):
    """Light rule on all four edges of a cell.

    The cover grid deliberately carries no table style — a style would draw
    borders on every band including the hero — so the project card asks for
    its own rules.
    """
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcPr.append(borders)
    for side in ("top", "left", "bottom", "right"):
        e = OxmlElement("w:%s" % side)
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:color"), colour)
        borders.append(e)


GREEN_RULE = "2F9E63"


def _cell_rules(cell, sides):
    """Draw rules on named edges only.

    _hairline boxes a cell on all four sides, which on the cover produced a
    grid of little boxes rather than a document table. sides is a mapping of
    edge name to (colour, size), e.g. {"bottom": ("D9E1E8", 4)}. Edges are
    written in the order OOXML expects.
    """
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(old)
    b = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        if side not in sides:
            continue
        colour, size = sides[side]
        e = OxmlElement(f"w:{side}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(size))
        e.set(qn("w:color"), colour)
        b.append(e)
    tcPr.append(b)


def _vcenter(cell):
    """Centre a cell's content vertically.

    Cells default to top alignment, so the logo and the accreditation strip
    sat on the same top edge despite being different heights, and read as
    two unrelated objects rather than one masthead.
    """
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:vAlign")):
        tcPr.remove(old)
    va = OxmlElement("w:vAlign")
    va.set(qn("w:val"), "center")
    tcPr.append(va)


def _row_height(row, cm, rule="atLeast"):
    """Give a row a floor height so a block of rows keeps an even rhythm."""
    trPr = row._tr.get_or_add_trPr()
    h = OxmlElement("w:trHeight")
    h.set(qn("w:val"), str(int(cm * 567)))
    h.set(qn("w:hRule"), rule)
    trPr.append(h)



def _fit_picture(paragraph, path, max_w_cm, max_h_cm, trim=True):
    """Insert a picture scaled to fit inside a box, trimming blank edges.

    Artwork from a certification body usually arrives centred on a large
    white canvas. Inserted at a fixed width that canvas becomes several
    centimetres of empty masthead — enough to push the cover onto a second
    page. Trimming the uniform border and then fitting to a box means the
    mark occupies the space it is given whatever shape the source file is.

    Falls back to a plain width-constrained insert if anything goes wrong;
    a slightly large logo is better than a build failure.
    """
    try:
        from PIL import Image, ImageChops
        im = Image.open(path)
        if trim:
            rgb = im.convert("RGB")
            corner = rgb.getpixel((0, 0))
            bg = Image.new("RGB", rgb.size, corner)
            box = ImageChops.difference(rgb, bg).getbbox()
            if box and box[2] - box[0] > 20 and box[3] - box[1] > 20:
                area_before = im.size[0] * im.size[1]
                cropped = im.crop(box)
                if cropped.size[0] * cropped.size[1] < area_before * 0.98:
                    import tempfile
                    fd, tmp = tempfile.mkstemp(suffix=".png")
                    os.close(fd)
                    cropped.save(tmp)
                    path, im = tmp, cropped
        w, h = im.size
        aspect = w / float(h) if h else 1.0
    except Exception:
        paragraph.add_run().add_picture(path, width=Cm(max_w_cm))
        return

    if max_w_cm / aspect <= max_h_cm:
        paragraph.add_run().add_picture(path, width=Cm(max_w_cm))
    else:
        paragraph.add_run().add_picture(path, height=Cm(max_h_cm))



def _make_watermark_png(src: str, dest: str, strength: float = 0.10) -> str:
    """Fade a logo into a watermark.

    Word can wash an image out through VML gain/blacklevel attributes, but
    LibreOffice ignores them and would print the mark at full strength across
    every page. Fading the pixels instead is renderer-independent.
    """
    from PIL import Image
    im = Image.open(src).convert("RGBA")
    white = Image.new("RGBA", im.size, (255, 255, 255, 255))
    faded = Image.blend(white, im, strength)
    faded.putalpha(im.getchannel("A").point(lambda a: int(a * strength * 3)))
    faded.save(dest)
    return dest


VML_NS = (
    'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
    'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
    'xmlns:v="urn:schemas-microsoft-com:vml" '
    'xmlns:o="urn:schemas-microsoft-com:office:office"'
)


def _watermark(section, image_path: str, width_pt: float = 230.0) -> bool:
    """Place a faded mark behind the text on every page of this section.

    A watermark is a VML shape anchored in the header with a negative
    z-index, which is what puts it behind the body text rather than over it.
    python-docx has no API for this, so the shape is built as raw XML.

    ``stroked="f"`` is not decoration. A #_x0000_t75 picture frame carries a
    default outline: Word does not draw it, LibreOffice does — so the DOCX
    looked clean while every page of the PDF carried a black rotated square
    around the mark. ``filled="f"`` suppresses the shape's own background for
    the same reason, leaving only the image itself.
    """
    try:
        from PIL import Image
        with Image.open(image_path) as probe:
            w, h = probe.size
        height_pt = width_pt * h / float(w)

        part = section.header.part
        rid, _ = part.get_or_add_image(image_path)

        para = section.header.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        xml = (
            '<w:r %s>'
            '<w:pict>'
            '<v:shape id="ecoreport_watermark" o:spid="_x0000_s2051" '
            'type="#_x0000_t75" stroked="f" filled="f" '
            'style="position:absolute;margin-left:0;margin-top:0;'
            'width:%.1fpt;height:%.1fpt;z-index:-251654144;'
            'mso-position-horizontal:center;'
            'mso-position-horizontal-relative:margin;'
            'mso-position-vertical:center;'
            'mso-position-vertical-relative:margin;'
            'rotation:-30" o:allowincell="f">'
            '<v:imagedata r:id="%s" o:title="watermark"/>'
            '</v:shape>'
            '</w:pict>'
            '</w:r>'
        # nsdecls only carries the WordprocessingML prefixes; v (VML) and o
        # (Office drawing) have to be declared by hand, and asking nsdecls for
        # them raised KeyError, which the except below swallowed. That is why
        # no report has ever carried the watermark.
        ) % (VML_NS, width_pt, height_pt, rid)
        para._p.append(parse_xml(xml))
        return True
    except Exception:  # noqa: BLE001
        import logging
        logging.getLogger(__name__).warning("watermark skipped", exc_info=True)
        return False



# A page holds roughly this many short table rows. Above it a table cannot
# be kept whole, so it is allowed to flow with its header repeating rather
# than forced into a layout that will not fit.
MAX_ROWS_KEPT_WHOLE = 24


def _typeset(doc) -> dict:
    """Apply publication-grade pagination controls across the whole document.

    Word will happily break a table row in half, drop a header row off the
    top of a continuation page, or strand a single line at the foot of a
    page. None of that is acceptable in a document going to a ministry under
    a stamp, and none of it can be fixed by editing content — it is governed
    by properties that have to be set on the paragraphs and rows themselves.

    Applied here as one pass over the finished template rather than at each
    of the thirty-odd call sites, so nothing can be missed when a new table
    or figure is added later.

    Returns a count of what was set, for the build log.
    """
    stats = {"rows_unsplit": 0, "header_rows_repeated": 0,
             "kept_with_table": 0, "figures_kept": 0, "widow_control": 0,
             "tables_kept_whole": 0}

    def _pPr(p):
        return p._p.get_or_add_pPr()

    def _flag(p, tag, attr=None):
        pPr = _pPr(p)
        if pPr.find(qn(tag)) is None:
            e = OxmlElement(tag)
            if attr:
                e.set(qn("w:val"), attr)
            pPr.append(e)
            return True
        return False

    # --- paragraphs -------------------------------------------------------
    body = list(doc.element.body)
    for i, el in enumerate(body):
        if not el.tag.endswith("}p"):
            continue
        para = Paragraph(el, doc)

        # never strand a single line at the top or bottom of a page
        if _flag(para, "w:widowControl"):
            stats["widow_control"] += 1

        text = (para.text or "").strip()

        # a paragraph carrying an image must not be parted from the caption
        # that follows it
        if "<pic:pic" in el.xml or "{{ fig_" in text or "{{ cover_hero" in text:
            para.paragraph_format.keep_with_next = True
            stats["figures_kept"] += 1
            continue

        # A lead-in line above a table stays with it. The look-ahead skips a
        # caption paragraph, because a table is nearly always introduced as
        # "...are summarized in Table 5.", then the caption, then the table.
        # Checking only the immediately following element meant the chain
        # broke at the lead-in: the caption and table moved to the next page
        # together while the sentence — and the section heading bound to it —
        # were left stranded at the foot of the previous one.
        nxt = body[i + 1] if i + 1 < len(body) else None
        after = body[i + 2] if i + 2 < len(body) else None
        leads_to_table = nxt is not None and nxt.tag.endswith("}tbl")
        if not leads_to_table and nxt is not None and nxt.tag.endswith("}p") \
                and after is not None and after.tag.endswith("}tbl"):
            nxt_style = Paragraph(nxt, doc).style
            leads_to_table = (nxt_style is not None
                              and (nxt_style.name or "") == "Caption")
        if leads_to_table and text:
            para.paragraph_format.keep_with_next = True
            stats["kept_with_table"] += 1

    # --- tables -----------------------------------------------------------
    for table in doc.tables:
        rows = table.rows
        for row in rows:
            trPr = row._tr.get_or_add_trPr()
            if trPr.find(qn("w:cantSplit")) is None:
                trPr.append(OxmlElement("w:cantSplit"))
                stats["rows_unsplit"] += 1

        # Hold the whole table on one page by binding every row to the one
        # after it. A table that will not fit on a page cannot be held
        # together — Word has to break it somewhere — so tables above the
        # threshold are left to flow with a repeating header instead of
        # being forced into an impossible layout.
        is_layout_grid = len(table.columns) >= 8
        if not is_layout_grid and 1 < len(rows) <= MAX_ROWS_KEPT_WHOLE:
            # Worked on the XML rather than through python-docx cells: a
            # vertically merged cell is reported as part of every row it
            # spans, and lxml creates element proxies on demand, so cells
            # cannot be compared by identity. In the tree a merged cell's
            # <w:tc> sits in the first row it covers and later rows carry
            # only a vMerge continuation, which is exactly the distinction
            # needed here.
            def _keep_next(tr, on):
                for tc in tr.findall(qn("w:tc")):
                    for p_el in tc.findall(qn("w:p")):
                        pPr = p_el.find(qn("w:pPr"))
                        if pPr is None:
                            if not on:
                                continue
                            pPr = OxmlElement("w:pPr")
                            p_el.insert(0, pPr)
                        existing = pPr.find(qn("w:keepNext"))
                        if on and existing is None:
                            pPr.append(OxmlElement("w:keepNext"))
                        elif not on and existing is not None:
                            pPr.remove(existing)

            for row in rows[:-1]:
                _keep_next(row._tr, True)
            # release the final row, or the table drags whatever follows it
            # onto its own page and leaves a gap behind
            _keep_next(rows[-1]._tr, False)
            stats["tables_kept_whole"] += 1

        # Repeat the header on continuation pages, but only for real data
        # tables. The cover is a twelve-column layout grid, not a data table:
        # marking its first row as a header made Word reprint the masthead at
        # the top of the following page.
        if len(rows) >= 4 and not is_layout_grid:
            trPr = rows[0]._tr.get_or_add_trPr()
            if trPr.find(qn("w:tblHeader")) is None:
                trPr.append(OxmlElement("w:tblHeader"))
                stats["header_rows_repeated"] += 1

    return stats


def build(out_path: str = OUT) -> str:
    doc = Document()
    st = doc.styles["Normal"]
    st.paragraph_format.line_spacing = 1.15
    st.paragraph_format.space_after = Pt(6)
    st.font.name = "Times New Roman"
    st.font.size = Pt(11)
    for lvl in range(1, 4):
        hs = doc.styles[f"Heading {lvl}"]
        hs.font.name = "Times New Roman"
        hs.font.color.rgb = {1: NAVY, 2: NAVY, 3: BLUE}[lvl]
        hs.font.size = Pt({1: 14, 2: 12, 3: 11}[lvl])
        hs.font.bold = True
    # accent rule under every level-1 heading
    h1 = doc.styles["Heading 1"]
    pPr = h1.element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "12")
    bot.set(qn("w:color"), BLUE_FILL)
    pBdr.append(bot)
    pPr.append(pBdr)

    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)  # A4
    sec.top_margin = sec.bottom_margin = Cm(2.2)
    sec.left_margin = sec.right_margin = Cm(2.2)

    # ---------------- COVER (full-bleed, no header/footer) ----------------
    sec.top_margin = sec.bottom_margin = Cm(0)
    sec.left_margin = sec.right_margin = Cm(0)

    def _pad(cell, top=0, bottom=0, left=0, right=0):
        tcPr = cell._tc.get_or_add_tcPr()
        mar = OxmlElement("w:tcMar")
        for tag, val in (("top", top), ("bottom", bottom),
                         ("left", left), ("right", right)):
            e = OxmlElement(f"w:{tag}")
            e.set(qn("w:w"), str(int(val * 567)))
            e.set(qn("w:type"), "dxa")
            mar.append(e)
        tcPr.append(mar)

    def _fill(cell, colour):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), colour)
        tcPr.append(shd)

    def _full_width(t, cols=1):
        tblPr = t._tbl.tblPr
        w = tblPr.find(qn("w:tblW"))
        if w is None:
            w = OxmlElement("w:tblW")
            tblPr.append(w)
        w.set(qn("w:w"), "5000")
        w.set(qn("w:type"), "pct")
        ind = OxmlElement("w:tblInd")
        ind.set(qn("w:w"), "0")
        ind.set(qn("w:type"), "dxa")
        tblPr.append(ind)
        cm = OxmlElement("w:tblCellMar")
        for tag in ("top", "left", "bottom", "right"):
            e = OxmlElement(f"w:{tag}")
            e.set(qn("w:w"), "0")
            e.set(qn("w:type"), "dxa")
            cm.append(e)
        tblPr.append(cm)
        lay = OxmlElement("w:tblLayout")
        lay.set(qn("w:type"), "fixed")
        tblPr.append(lay)
        grid = t._tbl.find(qn("w:tblGrid"))
        if grid is not None:
            each = int(21.0 * 567 / cols)
            for gc in grid.findall(qn("w:gridCol")):
                gc.set(qn("w:w"), str(each))
        return t

    def _txt(cell, text, size, bold=False, colour=None, align="left",
             before=0, after=0, italic=False, first=False):
        p = cell.paragraphs[0] if (first and not cell.paragraphs[0].text) \
            else cell.add_paragraph()
        p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT,
                       "center": WD_ALIGN_PARAGRAPH.CENTER,
                       "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        if colour is not None:
            r.font.color.rgb = colour
        return p

    # ------------------------------------------------------------------
    # Cover: one full-bleed table on a single twelve-column grid.
    #
    # Every band of the cover — masthead, hero, value propositions, project
    # card, footer — is a row of the same table, merged to whatever span it
    # needs. Nothing is nested. Nested tables carry their own width, and a
    # child whose fixed width disagrees with its parent cell overhangs the
    # margin or collapses; twelve columns divide cleanly by 2, 3, 4 and 6, so
    # every band lands on the same grid and the edges line up down the page.
    # The cover section already has zero margins, so the grid spans the full
    # 21 cm page.
    # ------------------------------------------------------------------
    # Fourteen columns: a narrow margin column at each edge, and twelve
    # content columns between them. Twelve divides by 2, 3, 4 and 6, so the
    # value propositions, the card and the footer all land on the same lines;
    # the margin columns keep those bands off the paper edge while the hero
    # and the footer band still bleed across the full page.
    COLS = 14
    MARGIN_W, CONTENT_W = 0.9, 19.2 / 12
    COL_WIDTHS = [MARGIN_W] + [CONTENT_W] * 12 + [MARGIN_W]
    C0, C1 = 1, 12                     # first and last content column
    ROW_MASTHEAD, ROW_HERO, ROW_ICON, ROW_PROP = 0, 1, 2, 3
    ROW_CARD = 4
    ROW_FOOT_NAME, ROW_FOOT_LINKS = 9, 10

    cov = doc.add_table(rows=11, cols=COLS)
    _full_width(cov, COLS)
    for row in cov.rows:
        for i, c in enumerate(row.cells):
            c.width = Cm(COL_WIDTHS[i])

    def _span(r, a, b):
        """Merge columns a..b of row r and return the resulting cell."""
        cell = cov.cell(r, a)
        if b > a:
            cell = cell.merge(cov.cell(r, b))
        for para in cell.paragraphs:
            para.paragraph_format.space_after = Pt(0)
        return cell

    def _left_rule(cell, colour="DCE5EE"):
        """Hairline on a cell's left edge — the dividers between the four
        value propositions in the approved cover."""
        tcPr = cell._tc.get_or_add_tcPr()
        borders = tcPr.find(qn("w:tcBorders"))
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            tcPr.append(borders)
        e = OxmlElement("w:left")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:color"), colour)
        borders.append(e)

    # --- row 0: masthead — BSA mark, then the accreditation marks ----------
    # A single strip image wins if present: certification bodies supply one
    # artwork with the marks and captions already set. The individual slots
    # remain for anyone who holds the marks separately, and a badge with no
    # file is skipped — the row never claims an accreditation the lab does
    # not hold.
    #
    # Both sides are vertically centred against each other, inset to the same
    # 1.25 cm axis the project card and the footer use, and the mark is given
    # more height than the accreditation strip so the masthead has an owner.
    # A green hairline closes the band, giving the page a top edge instead of
    # letting white run straight into the photograph.
    strip = os.path.join(ASSETS, "badges.png")
    badges = [] if os.path.exists(strip) else [
        (f, cap) for f, cap in ACCREDITATION_BADGES
        if os.path.exists(os.path.join(ASSETS, f))]

    INSET = 1.25
    BAND_PAD = 0.46

    logo_span = 4 if (badges or os.path.exists(strip)) else C1
    hl = _span(ROW_MASTHEAD, 0, logo_span)
    _pad(hl, BAND_PAD, BAND_PAD, INSET, 0.2)
    _vcenter(hl)
    try:
        _fit_picture(hl.paragraphs[0], os.path.join(ASSETS, "logo_left.png"),
                     max_w_cm=5.2, max_h_cm=1.85)
    except Exception:
        r = hl.paragraphs[0].add_run("BSA.lab")
        r.bold = True
        r.font.size = Pt(22)
        r.font.color.rgb = NAVY

    if os.path.exists(strip):
        sc = _span(ROW_MASTHEAD, logo_span + 1, COLS - 1)
        _pad(sc, BAND_PAD, BAND_PAD, 0.2, INSET)
        _vcenter(sc)
        sp = sc.paragraphs[0]
        sp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # held below the logo's height so the accreditation artwork supports
        # the mark rather than competing with it
        _fit_picture(sp, strip, max_w_cm=11.4, max_h_cm=1.5)
    elif badges:
        first = logo_span + 1
        each = max((COLS - first) // len(badges), 1)
        for n, (fname, caption) in enumerate(badges):
            a = first + n * each
            b = COLS - 1 if n == len(badges) - 1 else min(a + each - 1, COLS - 1)
            if a > COLS - 1:
                break
            cell = _span(ROW_MASTHEAD, a, b)
            # equal padding on every badge: unequal insets were what made
            # four marks of different proportions look scattered
            _pad(cell, BAND_PAD, BAND_PAD, 0.10, 0.10)
            _vcenter(cell)
            ip = cell.paragraphs[0]
            ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ip.paragraph_format.space_after = Pt(2)
            _fit_picture(ip, os.path.join(ASSETS, fname),
                         max_w_cm=2.8, max_h_cm=1.05)
            for i, line in enumerate(caption):
                cp = cell.add_paragraph()
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cp.paragraph_format.space_after = Pt(0)
                cp.paragraph_format.line_spacing = 1.0
                rr = cp.add_run(line)
                rr.font.size = Pt(5.5)
                rr.bold = (i == 0)
                rr.font.color.rgb = NAVY

    # one continuous rule across the full page width, including the margin
    # columns, so the band closes cleanly against the hero below it
    for c in cov.rows[ROW_MASTHEAD].cells:
        _cell_rules(c, {"bottom": (GREEN_RULE, 8)})

    # --- row 1: hero band, full bleed --------------------------------------
    hero_c = _span(ROW_HERO, 0, COLS - 1)
    _pad(hero_c, 0, 0, 0, 0)
    # Cell shading reaches the paper edge; an inline image does not — the
    # renderer insets the text flow by about 3 mm. Filling the cell with the
    # colour of the band's own left edge closes that strip, so the hero reads
    # as full bleed exactly like the footer below it.
    _fill(hero_c, HERO_EDGE_FILL)
    hp = hero_c.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.paragraph_format.space_after = Pt(0)
    # no indent of any kind, so the band bleeds evenly to both edges
    hp.paragraph_format.left_indent = Cm(0)
    hp.paragraph_format.right_indent = Cm(0)
    hp.add_run("{{ cover_hero }}")

    # --- rows 2-3: value propositions, three columns of the grid each ------
    props = [
        ("accurate", "ACCURATE", "Precision monitoring\nwith calibrated instruments"),
        ("reliable", "RELIABLE", "Data you can trust,\nanytime, anywhere"),
        ("compliant", "COMPLIANT", "Aligned with KSA NCEC\n& international standards"),
        ("sustainable", "SUSTAINABLE", "Supporting a cleaner\nand healthier future"),
    ]
    for i, (icon, title, blurb) in enumerate(props):
        a, b = C0 + i * 3, C0 + i * 3 + 2
        top = _span(ROW_ICON, a, b)
        bot = _span(ROW_PROP, a, b)
        # the grid's margin columns already hold the band off the page edge,
        # so the cells themselves only need even internal padding
        _pad(top, 0.30, 0.06, 0.2, 0.2)
        _pad(bot, 0.0, 0.26, 0.2, 0.2)
        if i:                       # no rule to the left of the first column
            _left_rule(top)
            _left_rule(bot)
        ip = top.paragraphs[0]
        ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ip.paragraph_format.space_after = Pt(3)
        try:
            ip.add_run().add_picture(
                os.path.join(ASSETS, f"icon_{icon}.png"), height=Cm(0.9))
        except Exception:
            pass
        _txt(bot, title, 9.5, bold=True, colour=GREEN_ACCENT, align="center",
             after=2, first=True)
        for line in blurb.split("\n"):
            _txt(bot, line, 8, colour=DARK, align="center", after=0)

    # --- rows 4-8: document control block -----------------------------------
    # Rebuilt as a document table rather than a card. The grey label fill was
    # opaque, so it punched holes in the watermark exactly as the zebra
    # striping did, and boxing every cell on four sides produced a grid of
    # small rectangles. Rules now run horizontally only, the labels sit on the
    # same 1.25 cm left axis as the footer, and a fixed row height gives the
    # block an even rhythm regardless of how long a value runs.
    rows_ = [("Client", "{{ client }}"),
             ("Monitoring station", "{{ site_name }}"),
             ("Survey period", "{{ monitoring_window_text }}"),
             ("Report number", "{{ report_number }}"),
             ("Revision and issue date",
              "{{ revision }}     \u00b7     {{ reporting_date }}")]
    for i, (k, v) in enumerate(rows_):
        r = ROW_CARD + i
        kc = _span(r, C0, C0 + 2)          # three columns — 4.8 cm
        vc = _span(r, C0 + 3, C1)          # nine columns — 14.4 cm
        _pad(kc, 0.16, 0.16, 0.35, 0.25)
        _pad(vc, 0.16, 0.16, 0.0, 0.35)
        _vcenter(kc)
        _vcenter(vc)
        _row_height(cov.rows[r], 0.76)
        rules = {"bottom": (RULE_GREY, 4)}
        if i == 0:
            rules["top"] = (GREEN_RULE, 8)
        _cell_rules(kc, rules)
        _cell_rules(vc, rules)
        # label quiet and small, value dark and larger: the weight difference
        # carries the hierarchy, so the labels no longer need to shout in caps
        _txt(kc, k, 8.5, colour=MUTED_GREY, first=True)
        _txt(vc, v, 10, colour=DARK, first=True)

    # The approved concept carries no prepared-by block on the cover, and
    # the same detail is already set out on the document control page. Its
    # removal is also what lets the cover close on a single page now that the
    # hero band is taller.

    # --- rows 10-11: navy contact footer ------------------------------------
    fn = _span(ROW_FOOT_NAME, 0, COLS - 1)
    _fill(fn, NAVY_FILL)
    _pad(fn, 0.34, 0.10, 1.25, 1.25)
    _txt(fn, "{{ provider_legal_name }}", 10.5, bold=True, colour=WHITE,
         first=True)

    # uneven spans: the address is the longest item and was wrapping to a
    # second line once the outer columns were inset for print safety
    items = [("{{ provider_website }}", C0, C0 + 2),
             ("{{ provider_email }}", C0 + 3, C0 + 5),
             ("{{ provider_tel }}", C0 + 6, C0 + 7),
             ("{{ provider_address }}", C0 + 8, C1)]
    for edge in (0, COLS - 1):          # keep the navy bleeding to the edge
        _fill(cov.cell(ROW_FOOT_LINKS, edge), NAVY_FILL)
    for i, (item, a, b) in enumerate(items):
        c = _span(ROW_FOOT_LINKS, a, b)
        _fill(c, NAVY_FILL)
        # the outer columns are inset to the same safe margin as the rest of
        # the cover: on a full-bleed page the last contact item was ending
        # 2.6 mm from the trim edge, close enough to be cut off in print
        _pad(c, 0.04, 0.34, 0.15, 0.15)
        _txt(c, item, 8.5, colour=RGBColor(0xC5, 0xDA, 0xEC),
             align="center", first=True)

    sec2 = doc.add_section(WD_SECTION.NEW_PAGE)
    sec2.page_width, sec2.page_height = Cm(21.0), Cm(29.7)
    sec2.top_margin = Cm(2.3)
    sec2.bottom_margin = Cm(2.0)
    sec2.left_margin = sec2.right_margin = Cm(2.3)
    sec2.top_margin = Cm(3.2)
    _header_footer(sec2)

    # --- Document control page
    _p(doc, "Ambient Air Quality Monitoring Report", size=16, bold=True,
       align="center",
       space_after=2)
    _p(doc, "for", size=12, align="center", space_after=2)
    _p(doc, "{{ project_name }}", size=15, bold=True, align="center", space_after=10)
    _p(doc, "Prepared by", size=12, align="center", space_after=2)
    _p(doc, "{{ provider }}", size=13, bold=True, align="center", space_after=14)

    meta = doc.add_table(rows=4, cols=2)
    meta.style = "Table Grid"
    _polish(meta)
    for i, (k, v) in enumerate([
        ("Project", "{{ project_name }}"),
        ("Document Title", "Ambient Air Quality Monitoring Report"),
        ("Client", "{{ client }}"),
        ("Report Number", "{{ report_number }}"),
    ]):
        _cell_text(meta.cell(i, 0), k, bold=True, size=10)
        _shade(meta.cell(i, 0))
        _cell_text(meta.cell(i, 1), v, size=10)
    _p(doc, space_after=8)
    _p(doc, "Revision History", size=11, bold=True, color=NAVY, space_after=4)
    rev = doc.add_table(rows=2, cols=5)
    rev.style = "Table Grid"
    _polish(rev)
    for j, h in enumerate(["REV", "REPORTING DATE", "PREPARED BY",
                           "PROJECT SUPERVISION", "STATUS"]):
        _cell_text(rev.cell(0, j), h, bold=True, size=9, align="center")
        _shade(rev.cell(0, j))
    for j, v in enumerate(["{{ revision }}", "{{ reporting_date }}",
                           "{{ prepared_by }}", "{{ project_supervision }}",
                           "{{ document_status }}"]):
        _cell_text(rev.cell(1, j), v, size=9.5, align="center")
    _p(doc, space_after=10)

    _p(doc, "Accreditation", size=11, bold=True, color=NAVY, space_after=4)
    _p(doc, "{{ provider_legal_name }} is an environmental laboratory approved "
            "by the National Center for Environmental Compliance (NCEC) of the "
            "Kingdom of Saudi Arabia for ambient air quality monitoring. "
            "Monitoring, quality assurance and data validation are undertaken "
            "in accordance with the applicable USEPA reference methods and "
            "data-handling guidance.", align="justify", size=10)
    _p(doc, "Calibration of the gaseous analysers is performed against "
            "internationally traceable reference standards; the expanded "
            "uncertainty stated on the certificates is ±2% (k = 2, "
            "approximately 95% confidence).", align="justify", size=10)
    doc.add_page_break()

    # --- TOC / LoF / LoT
    _p(doc, "Table of Contents", size=14, bold=True, space_after=8)
    toc_p = doc.add_paragraph()
    _field(toc_p, ' TOC \\o "1-3" \\h \\z \\u ')
    doc.add_page_break()
    _p(doc, "List of Figures", size=14, bold=True, space_after=8)
    lof_p = doc.add_paragraph()
    _field(lof_p, ' TOC \\h \\z \\c "Figure" ')
    _p(doc, "List of Tables", size=14, bold=True, space_after=8)
    lot_p = doc.add_paragraph()
    _field(lot_p, ' TOC \\h \\z \\c "Table" ')
    doc.add_page_break()

    # --- Definitions & Abbreviations
    _p(doc, "Definitions & Abbreviations", size=14, bold=True, space_after=8)
    defs = [
        ("°C", "Degrees Celsius"),
        ("µg/m³", "Micrograms per cubic meter at standard temperature and pressure "
                  "(25°C and 101.3 kPa)"),
        ("W/m²", "Watt per square meter"),
        ("AAQMS", "Ambient Air Quality Monitoring Station"),
        ("CO", "Carbon Monoxide"),
        ("Deg.", "Degrees (True North)"),
        ("H₂S", "Hydrogen sulfide"),
        ("m/s", "Meters per second"),
        ("NO", "Nitric oxide"),
        ("NO2", "Nitrogen dioxide"),
        ("NOx", "Oxides of Nitrogen"),
        ("O3", "Ozone"),
        ("hPa", "Hecto Pascal"),
        ("PM10", "Particulate less than 10 microns in equivalent aerodynamic diameter"),
        ("PM2.5", "Particulate less than 2.5 microns in equivalent aerodynamic diameter"),
        ("ppb", "Parts per billion"),
        ("ppm", "Parts per million"),
        ("RH", "Relative Humidity"),
        ("SO₂", "Sulfur dioxide"),
        ("WD", "Wind Direction"),
        ("WS", "Wind Speed"),
    ]
    dt = doc.add_table(rows=len(defs) + 1, cols=2)
    dt.style = "Table Grid"
    _polish(dt, header_rows=1)
    for j, h in enumerate(["ABBREVIATION", "DEFINITION"]):
        _cell_text(dt.cell(0, j), h, bold=True, size=9,
                   align="center" if j == 0 else "left")
        _shade(dt.cell(0, j))
    for i, (a, b) in enumerate(defs, start=1):
        dt.cell(i, 0).width = Cm(3.2)
        dt.cell(i, 1).width = Cm(13.3)
        _cell_text(dt.cell(i, 0), a, bold=True, size=9.5, align="center")
        _cell_text(dt.cell(i, 1), b, size=9.5)
    doc.add_page_break()

    # --- Executive Summary
    _heading(doc, "Executive Summary", 1)
    _p(doc, "{{ provider }} was commissioned by {{ client }} to conduct ambient air "
            "quality monitoring at {{ project_name }}. Continuous ambient air "
            "monitoring was conducted at one location for a period of "
            "{{ monitoring_period_text }} by the air quality monitoring station "
            "(AQMS). The Air Quality Monitoring started on {{ monitoring_start_date }}. "
            "This report presents a summary of the validated data that was obtained "
            "for the period of {{ monitoring_window_text }}. The ambient air quality "
            "monitoring station (AAQMS) was equipped to measure standard air "
            "pollutants and meteorological parameters as listed below.",
       align="justify")
    lst = doc.add_table(rows=8, cols=2)
    heads = ["Air Pollutants Monitored", "Meteorological Parameters Monitored"]
    for j, h in enumerate(heads):
        _cell_text(lst.cell(0, j), h, bold=True, size=10)
        _shade(lst.cell(0, j))
    pol_names = ["Oxides of Nitrogen (NO2)", "Sulphur Dioxide (SO₂)",
                 "Carbon Monoxide (CO)", "Ozone (O3)", "Hydrogen Sulfide (H₂S)",
                 "Particulates Matter - PM10", "Particulates Matter - PM2.5"]
    met_names = ["Wind Speed", "Wind Direction", "Air Temperature",
                 "Relative Humidity", "Barometric Pressure", "", ""]
    for i in range(7):
        _cell_text(lst.cell(i + 1, 0), pol_names[i], size=10)
        _cell_text(lst.cell(i + 1, 1), met_names[i], size=10)
    _p(doc)
    _p(doc, "{{ capture_sentence }} For QA/QC checks, in accordance with the relevant "
            "United States Environmental Protection Agency (EPA) methods for each "
            "parameter, were carried out within the required schedule. USEPA data "
            "handling guidelines were followed in collecting, verifying, and "
            "validating continuous ambient air quality and meteorological monitoring "
            "data in this report.", align="justify")
    _p(doc, space_after=6)
    _callout(doc, "KEY FINDING", "{{ headline_finding }}")
    _p(doc, space_after=8)
    _p(doc, "{%p if site_geometry_text %}", size=1, space_after=0)
    _callout(doc, "MONITORING LOCATION", "{{ site_geometry_text }}",
             fill="F1F6F1", edge="2F7D32")
    _p(doc, space_after=8)
    _p(doc, "{%p endif %}", size=1, space_after=0)
    _caption(doc, "Table", "Percent of Data Captured for all Parameters.")
    cap = doc.add_table(rows=4, cols=5)
    cap.style = "Table Grid"
    _polish(cap)
    for j, h in enumerate(["Parameters", "Total hours in monitoring period",
                           "Total available hours in monitoring period",
                           "Exception hours", "AAQMS 1-Hour's data capture %"]):
        _cell_text(cap.cell(0, j), h, bold=True, size=9, align="center")
        _shade(cap.cell(0, j))
    _tr_tag_row(cap, 1, "{%tr for r in capture_rows %}")
    row = cap.rows[2]
    _cell_text(row.cells[0], "{{ r.name }}", size=9)
    _cell_text(row.cells[1], "{{ r.total }}", size=9, align="center")
    _cell_text(row.cells[2], "{{ r.available }}", size=9, align="center")
    _cell_text(row.cells[3], "{{ r.exception }}", size=9, align="center")
    _cell_text(row.cells[4], "{{ r.capture }}", size=9, align="center")
    _tr_tag_row(cap, 3, "{%tr endfor %}")
    doc.add_page_break()

    # --- 1. Introduction
    _heading(doc, "1. Introduction", 1)
    _p(doc, "An ambient air quality monitoring survey was conducted for the "
            "{{ project_name }}. {{ provider }} ({{ provider_short }} - is an "
            "environmental laboratory in field of air quality monitoring approved by "
            "NCEC) installed the AQMS at one location in the proposed location as per "
            "client request and environmental judgment. {{ provider_short }} stations "
            "Lab was retained by {{ client }} and {{ provider_short }} was responsible "
            "for the operation and maintenance of the AAQMS as well as the validation "
            "of the data recorded. This report presents the data collected at the "
            "project site for the period of {{ monitoring_window_text }}.",
       align="justify")
    _p(doc, "This report summarizes the results obtained from the ambient air quality "
            "survey and field observations for any exceedances of the NCEC air "
            "quality standard. The ambient air quality standards used to identify "
            "pollution include the national standards set out in the Implementing "
            "Regulations for Air Quality of the Environmental Law issued by Royal "
            "Decree No. (M/165) as of 19/11/1441 AH. Graphical representations of "
            "the monitoring results within the context of the relevant limit values "
            "are also provided.", align="justify")
    _p(doc, "The following air pollutants were measured at each point:")
    for b in ["Particulate matter with aerodynamic diameters less than 10 microns (PM10),",
              "Particulate matter with aerodynamic diameters less than 2.5 microns (PM2.5),",
              "Sulfur dioxide (SO₂).", "Hydrogen Sulfide (H₂S).",
              "Oxides of Nitrogen (NO2).", "Ozone (O3).", "Carbon Monoxide (CO)."]:
        doc.add_paragraph(b, style="List Bullet")
    _p(doc, "In addition, meteorological data for wind speed, wind direction, ambient "
            "temperature, relative humidity, and barometric pressure were also "
            "measured.", align="justify")

    # --- 2. Monitoring and Data Collection
    _heading(doc, "2. Monitoring and Data Collection", 1)
    _heading(doc, "2.1 Site Details", 2)
    _p(doc, "The location of the ambient air quality monitoring station as shown in "
            "Table 2 and Figure 1. The Monitoring location was chosen by "
            "({{ client }}) and was intended to provide background or baseline data "
            "for the site. Inlet manifold length {{ inlet_height_m }} meters from "
            "ground level.", align="justify")
    _caption(doc, "Table", "Location of Ambient Air Quality Monitoring Stations")
    loc = doc.add_table(rows=2, cols=2)
    loc.style = "Table Grid"
    _polish(loc)
    _cell_text(loc.cell(0, 0), "Site Name", bold=True, size=10, align="center")
    _shade(loc.cell(0, 0))
    _cell_text(loc.cell(0, 1), "Geographical Coordinates", bold=True, size=10,
               align="center")
    _shade(loc.cell(0, 1))
    _cell_text(loc.cell(1, 0), "{{ site_name }}", size=10, align="center")
    _cell_text(loc.cell(1, 1), "N {{ latitude }}   E {{ longitude }}", size=10,
               align="center")
    # Each caption sits inside its figure's own condition. Left outside, the
    # caption printed whether or not the figure did — a report with no site
    # map carried "Figure 1 — Location of the Ambient Air quality monitor"
    # above empty space, and the SEQ field still consumed a figure number, so
    # the List of Figures described a figure the reader could not find.
    _p(doc, "{%p if fig_site_map %}", size=1, space_after=0)
    _p(doc, "{{ fig_site_map }}", align="center")
    _caption(doc, "Figure", "Location of the Ambient Air quality monitor")
    _p(doc, "{%p endif %}", size=1, space_after=0)
    grid = doc.add_table(rows=3, cols=2)
    grid.alignment = WD_TABLE_ALIGNMENT.CENTER
    _tr_tag_row(grid, 0, "{%tr for row in site_photo_rows %}")
    mid = grid.rows[1]
    for k, cell in enumerate(mid.cells):
        cp = cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.add_run("{{ row[%d] }}" % k)
    _tr_tag_row(grid, 2, "{%tr endfor %}")
    _p(doc, "{%p if fig_site_photo %}", size=1, space_after=0)
    _p(doc, "{{ fig_site_photo }}", align="center")
    _caption(doc, "Figure",
             "Location of the Ambient Air quality monitoring stations in the site")
    _p(doc, "{%p endif %}", size=1, space_after=0)

    _heading(doc, "2.2 Monitoring Methodology", 2)
    _p(doc, "Monitoring methodology and Reference Measurement Principle and "
            "Calibration Procedures for the Measurement of Ambient Air Quality "
            "Pollutants are summarized in Table 3.", align="justify")
    _caption(doc, "Table", "Reference Measurement Principle and equivalent reference method")
    m = doc.add_table(rows=7, cols=5)
    m.style = "Table Grid"
    _polish(m)
    for j, h in enumerate(["Measured Parameter", "Data Collection Methods Used",
                           "Description of the Method", "Method Type",
                           "Automated Equivalent Reference Method ID"]):
        _cell_text(m.cell(0, j), h, bold=True, size=9, align="center")
        _shade(m.cell(0, j))
    method_rows = [
        ("Sulfur Dioxide (SO2)", "40 CFR Appendix A-1 to Part 50",
         "Reference Measurement Principle and Calibration Procedure for the "
         "Measurement of Sulfur Dioxide in the Atmosphere (Ultraviolet Fluorescence "
         "Method)", "FRM", "EQSA-0486-060"),
        ("Nitrogen Oxides (NO - NO2 - NOx)",
         "Code of Federal Regulations, Title 40, Part 50, Appendix F",
         "Measurement Principle and Calibration Procedure for the Measurement of "
         "Nitrogen Dioxide in the Atmosphere (Gas Phase Chemiluminescence)",
         "FRM", "RFNA-1289-074"),
        ("Ozone (O3)", "USEPA Code of Federal Regulations (Title 40, Part 50, "
         "Appendix D)", "Reference Measurement Principle and Calibration Procedure "
         "for the Measurement of Ozone in the Atmosphere (UV light absorption "
         "Method)", "FRM", "EQOA-0880-047"),
        ("Carbon Monoxide (CO)", "40 CFR Appendix C to Part 50",
         "Measurement Principle and Calibration Procedure for the Measurement of "
         "Carbon Monoxide in the Atmosphere (Non-Dispersive Infrared Photometry)",
         "FRM", "RFCA-0981-054"),
        ("Hydrogen Sulfide (H2S)", "Described below in detail", "NA", "NA", "NA"),
        ("Particulate Matter (PM10 & PM2.5)", "40 CFR 50 Appendix L",
         "Reference Method for the Determination of Fine Particulate Matter",
         "FRM", "EQPM-0609-183"),
    ]
    for i, r in enumerate(method_rows):
        for j, v in enumerate(r):
            _cell_text(m.cell(i + 1, j), v, size=8)
    _p(doc, "* FRM: Federal Reference Method", size=9, italic=True)

    _heading(doc, "2.3 Monitored Parameters", 2)
    _p(doc, "Table 4 below shows the details of parameters monitored and the "
            "instruments used at the monitoring station. Note that the "
            "meteorological instruments (wind Speed, wind direction, Temp, and RH) "
            "sensors installed at 10 m above ground level. Barometric pressure "
            "installed at 3 m above ground level.", align="justify")
    _caption(doc, "Table", "Parameters Measured at Air Quality Monitoring Stations")
    inst = doc.add_table(rows=4, cols=3)
    inst.style = "Table Grid"
    _polish(inst)
    for j, h in enumerate(["PARAMETERS MEASURED", "SN",
                           "INSTRUMENT AND MEASUREMENTS TECHNIQUES"]):
        _cell_text(inst.cell(0, j), h, bold=True, size=9, align="center")
        _shade(inst.cell(0, j))
    _tr_tag_row(inst, 1, "{%tr for i in instruments %}")
    row = inst.rows[2]
    _cell_text(row.cells[0], "{{ i.parameter }}", size=9)
    _cell_text(row.cells[1], "{{ i.sn }}", size=9, align="center")
    _cell_text(row.cells[2], "{{ i.technique }}", size=9)
    _tr_tag_row(inst, 3, "{%tr endfor %}")
    _p(doc, "SN: serial number", size=9, italic=True)

    _heading(doc, "2.4 Data Collection Methods", 2)
    _heading(doc, "2.4.1 Compliance with Standards", 3)
    _p(doc, "The instruments used in the stations were approved by the US EPA and "
            "all procedures for ambient air monitoring are in accordance with US "
            "EPA procedures.", align="justify")
    _heading(doc, "2.4.2 Data Acquisition", 3)
    _p(doc, "Data acquisition was performed by using a PC based situated at each of "
            "the monitoring sites. Data was logged every 15 seconds and stored every "
            "1 minute. The data was backed up and downloaded every day and subjected "
            "to a rigorous program of quality checking.", align="justify")

    _heading(doc, "2.5 Data Validation and Reporting", 2)
    _heading(doc, "2.5.1 Validation", 3)
    _p(doc, "The purpose of the data validation section is to specify the guidelines "
            "that have been taken into consideration whenever possible in "
            "collecting, verifying, and validating continuous ambient air quality "
            "and meteorological monitoring data in this project.", align="justify")
    _p(doc, "{{ provider_short }} maintains one distinct database containing "
            "non-validated and validated data respectively. The validated database "
            "is created by duplicating the non-validated database and then flagging "
            "data affected by instrument faults, calibrations, and other maintenance "
            "activities. Invalid data is removed with the support of valid reason "
            "e.g., supported by maintenance notes, calibration sheets etc. as "
            "documented in exception tables.", align="justify")
    _p(doc, "Validation is performed by a trained data analyst. All data is checked "
            "and reviewed by the {{ provider_short }} section head. Graphs and "
            "reports are generated based on the validated hourly data and daily "
            "data.", align="justify")
    _heading(doc, "2.5.2 Data Validation Levels", 3)
    for lead, body in [
        ("Level 0", "Routine checks – Field and laboratory operations, data "
         "processing, reporting conducted in accordance with Standard Operating "
         "Procedures (SOPs) – Proper data file identification; review of unusual "
         "events, field data sheets, and result reports; instrument performance "
         "checks."),
        ("Level I", "Internal consistency tests – Identify values that appear "
         "atypical when compared to the values of the entire dataset."),
        ("Level II/III", "External consistency tests – Identify values in the data "
         "that appear atypical when compared to other datasets – Continued "
         "evaluation of the data as part of the data interpretation process."),
    ]:
        p = _p(doc, "", align="justify")
        r = p.add_run(lead)
        r.bold = True
        p.add_run(" — " + body)
    _p(doc, "Outliers are values that lie outside most of the other values in a set "
            "of data. Outliers treated as valid/suspect until proven invalid. The "
            "first assumption upon finding a measurement that is inconsistent with "
            "physical expectations is that the unusual value is due to a measurement "
            "error. If upon tracing the path of the measurement, nothing unusual is "
            "found, the value can be assumed to be a valid result of an "
            "environmental cause.", align="justify")
    _heading(doc, "2.5.3 Considerations in Evaluating AAQ Data", 3)
    for b in ["Levels of other pollutants", "Time of day/year",
              "Observations at other sites", "Audits and inter-laboratory comparisons",
              "Instrument performance history", "Calibration drift",
              "Site characteristics", "Meteorology", "Exceptional events"]:
        doc.add_paragraph(b, style="List Bullet")
    _p(doc, "Visual Data Review: Time Series are visually inspected for the below:")
    for b in ["Jumps, dips.", "Periodicity of peaks", "Calibration gas, carryover",
              "Expected diurnal pattern.", "Expected relationships.",
              "High concentrations of less abundant species or low concentrations "
              "of more abundant species"]:
        doc.add_paragraph(b, style="List Bullet")
    _p(doc, "Data is put into perspective of:")
    for b in ["Local, regional, or national averages", "Trends over time",
              "Comparison to nearby sites, similar areas", "Detection limits"]:
        doc.add_paragraph(b, style="List Bullet")
    _heading(doc, "2.5.4 Reporting", 3)
    _p(doc, "The reported data is provided in a Microsoft Excel spreadsheet. The "
            "data contained in these reports is based on Kingdom of Saudi Arabia "
            "(KSA) time.", align="justify")

    # --- 3. Standards
    _heading(doc, "3. Ambient Air Quality Standards", 1)
    _p(doc, "The air quality monitoring station's data has been compared to the "
            "NCEC's 2020 ambient air quality standards. Table 5 summarizes the "
            "NCEC's 2020 ambient air quality standards for the pollutants studied. "
            "Note that CO and O3 averages are rolling averages whereas all other "
            "averages are fixed averages.", align="justify")
    _caption(doc, "Table", "KSA NCEC Ambient Air Quality Standards")
    ncec = doc.add_table(rows=4, cols=5)
    ncec.style = "Table Grid"
    _polish(ncec)
    for j, h in enumerate(["Parameter", "Time Period", "Exceedance Level", "Units",
                           "Number of Allowable Exceedances"]):
        _cell_text(ncec.cell(0, j), h, bold=True, size=9, align="center")
        _shade(ncec.cell(0, j))
    _tr_tag_row(ncec, 1, "{%tr for l in ncec_rows %}")
    row = ncec.rows[2]
    _cell_text(row.cells[0], "{{ l.pollutant }}", size=9, align="center")
    _cell_text(row.cells[1], "{{ l.period }}", size=9, align="center")
    _cell_text(row.cells[2], "{{ l.limit }}", size=9, align="center")
    _cell_text(row.cells[3], "µg/m³", size=9, align="center")
    _cell_text(row.cells[4], "{{ l.allowance }}", size=9, align="center")
    _tr_tag_row(ncec, 3, "{%tr endfor %}")

    # --- 4. Calibrations and Maintenance
    _heading(doc, "4. Calibrations and Maintenance", 1)
    _heading(doc, "4.1 Maintenance", 2)
    _p(doc, "Regular maintenance was carried out by {{ provider_short }}. Sample "
            "inlets of all analyzers were cleaned before installation and on a "
            "weekly basis. Filters of gas analyzers were changed before the start "
            "of monitoring. Meteorological sensors cleaned before installation. In "
            "addition, the regular maintenance program conducted as per standard or "
            "operation manual such as: meteorological sensors should be cleaned "
            "after rain because the dust particles stick to the meteorological "
            "sensors. After every sandstorm, inlets of all analyzers as well as "
            "manifolds of the gas analyzers should be cleaned to protect analyzers "
            "from particulate matter.", align="justify")
    _heading(doc, "4.2 Calibration", 2)
    _p(doc, "Multipoint calibration of the instruments monitoring gaseous pollutants "
            "was conducted one time before the monitoring period. The instruments "
            "were calibrated in full accordance with the manufacturer's "
            "recommendations and conform to the requirements of USEPA. Calibrations "
            "were performed at the monitoring site by allowing the analyzer to "
            "sample a gaseous standard containing a known pollutant concentration. "
            "Calibration data were recorded by the same data acquisitions system. "
            "Mast wind direction system was oriented at South direction as per "
            "manufacture operation manual.", align="justify")

    # --- 5. Results and discussion
    _heading(doc, "5. Results and discussion", 1)
    _heading(doc, "5.1 Air Quality Summary", 2)
    _p(doc, "Table 6 summarises the compliance status of every monitored "
            "pollutant against the applicable NCEC 2020 limits. Detailed "
            "results for each pollutant follow in the sections below.",
       align="justify")
    _caption(doc, "Table", "Compliance Summary Matrix")
    cs = doc.add_table(rows=4, cols=7)
    cs.style = "Table Grid"
    _polish(cs)
    for j, h in enumerate(["POLLUTANT", "HOURLY MAX", "8-HR MAX", "DAILY AVG",
                           "APPLICABLE LIMIT", "% OF LIMIT", "STATUS"]):
        _cell_text(cs.cell(0, j), h, bold=True, size=8.5, align="center")
        _shade(cs.cell(0, j))
    # Explicit widths proportional to content. With seven equal columns Word
    # has to break "APPLICABLE" mid-word to fit it, which is why the header
    # printed as "APPLICABL / E LIMIT".
    for j, w in enumerate((2.6, 2.2, 2.0, 2.1, 3.1, 2.0, 2.4)):
        for r_ in cs.rows:
            r_.cells[j].width = Cm(w)
    _no_word_split(cs)
    _tr_tag_row(cs, 1, "{%tr for r in compliance_rows %}")
    row = cs.rows[2]
    _cell_text(row.cells[0], "{{ r.pollutant }}", size=9, bold=True)
    for k, key in enumerate(["hourly_max", "rolling_8h", "daily_avg", "limit",
                             "pct_of_limit"], start=1):
        _cell_text(row.cells[k], "{{ r.%s }}" % key, size=9, align="center")
    # Status is supplied as a RichText value so its colour is set with the
    # verdict itself; inline Jinja colour tags would break the row loop.
    _cell_text(cs.cell(2, 6), "{{ r.verdict }}", size=8.5, align="center",
               bold=True)
    _tr_tag_row(cs, 3, "{%tr endfor %}")

    lg = doc.add_paragraph()
    lg.paragraph_format.space_before = Pt(3)
    lg.paragraph_format.space_after = Pt(2)
    for label, colour in (("\u25A0 COMPLIANT", OK_GREEN),
                          ("\u25A0 SEE NOTE \u2014 above limit, allowance not yet reached",
                           WARN_AMBER),
                          ("\u25A0 EXCEEDANCE", BAD_RED)):
        rr = lg.add_run(label + "     ")
        rr.font.size = Pt(7.5)
        rr.font.color.rgb = colour
        rr.bold = True
    _p(doc, "Values are in \u00b5g/m\u00b3. The applicable limit shown is the "
            "NCEC averaging period the data comes closest to; all periods are "
            "assessed in the detailed tables that follow.", size=8.5,
       italic=True, color=MUTED_GREY)
    _p(doc, space_after=6)
    _p(doc, "Tables 6 to 13 compare monitoring results for AAQMS for the period of "
            "{{ monitoring_window_text }} on the site. The results are explained as "
            "follows:", align="justify")
    _p(doc, "- Air Pollutants Monitored results", bold=True)

    # 5.1.1 SO2
    _heading(doc, "5.1.1 Sulphur Dioxide (SO2)", 3)
    _p(doc, "Levels of Sulphur dioxide (SO2) in ambient air are typically directly "
            "related to the concentration of Sulphur in fuel and the quantity of "
            "fuel being combusted. Upon combustion, approximately 98% of the "
            "Sulphur in the fuel will oxidize to form SO2, with the remaining 2% "
            "producing Sulphur trioxide (SO3). The emitted SO2 can also further "
            "oxidize to SO3 and react with water to produce acid rain in the form "
            "of sulphury acid (H2SO4). Short-term exposures to SO2 have shown "
            "adverse respiratory effects including bronchoconstriction and "
            "increased asthma symptoms.", align="justify")
    _p(doc, "{{ so2.narrative }}", align="justify")
    _caption(doc, "Table", "Summary of SO₂ Results")
    _summary_table(doc, [
        ("Percentage data capture (Hourly Values)", "{{ so2.capture }}"),
        ("Hourly Maximum (ug/m³)", "{{ so2.h_max }}"),
        ("Hourly Minimum (ug/m³)", "{{ so2.h_min }}"),
        ("Daily average (ug/m³)", "{{ so2.daily_avg }}"),
        ("Hourly value > {{ so2.limit_1h }} (ug/m³)", "{{ so2.exceed_1h }}"),
    ], [("1 Hour", "{{ so2.limit_1h }}"), ("24 Hour", "{{ so2.limit_24h }}")])
    _p(doc, "{{ so2.footnote }}", size=9, italic=True)
    _verdict_box(doc, "{{ so2.verdict_line }}")
    _figure(doc, "fig_so2", "SO2 Hourly Concentration at the location.")

    # 5.1.2 NO/NO2/NOx
    _heading(doc, "5.1.2 Oxides of Nitrogen (NO, NO2, NOx)", 3)
    _p(doc, "In a combustion process, NOx is produced through three mechanisms, "
            "namely thermal NOx, fuel NOx and prompt NOx. Thermal NOx is the primary "
            "source of NOx and is formed as a high temperature dissociation and "
            "subsequent reaction of nitrogen (N2) and oxygen (O2). NO2 is the "
            "primary component of concern in NOx emissions. Generally, up to 10% of "
            "the NOx emitted from the combustion of fuel is emitted as NO2. The "
            "remainder is emitted as NO, which is subsequently converted to NO2 in "
            "reactions with various oxidants and ozone as the plume is transported "
            "downwind from the source. NO2 is a reddish-brown gas with a pungent "
            "odor, which upon reaction with other atmospheric compounds, becomes a "
            "major contributor to smog, acid rain, inhalable particulates and "
            "reduced visibility. At significant levels and exposure, inhalation may "
            "result in irritation and burning to the skin and eyes, nose, and "
            "throat. Prolonged exposure may result in permanent lung damage.",
       align="justify")
    _p(doc, "{{ nox_group.narrative }}", align="justify")
    _caption(doc, "Table", "Summary of (NO, NO2, NOx) Results.")
    nx = doc.add_table(rows=16, cols=3)
    nx.style = "Table Grid"
    _polish(nx)
    hdr = _merge_row(nx, 0, 0, 1)
    _cell_text(hdr, "NO₂ Concentration at sampling point", bold=True, size=10)
    _shade(hdr)
    _cell_text(nx.cell(0, 2), "NCEC Exceedance Level µg/m³ — 1 Hour", bold=True,
               size=9, align="center")
    _shade(nx.cell(0, 2))
    no2_rows = [
        ("Percentage data capture (Hourly Values)", "{{ no2.capture }}"),
        ("Hourly Maximum (ug/m³)", "{{ no2.h_max }}"),
        ("Hourly Minimum (ug/m³)", "{{ no2.h_min }}"),
        ("Hourly value > {{ no2.limit_1h }} (ug/m³)", "{{ no2.exceed_1h }}"),
        ("Daily average (ug/m³)", "{{ no2.daily_avg }}"),
    ]
    for i, (a, b) in enumerate(no2_rows):
        _cell_text(nx.cell(1 + i, 0), a, size=10)
        _cell_text(nx.cell(1 + i, 1), b, size=10, align="center")
    lim_cell = nx.cell(1, 2).merge(nx.cell(5, 2))
    _cell_text(lim_cell, "{{ no2.limit_1h }}", size=10, align="center")
    hdr2 = _merge_row(nx, 6, 0, 1)
    _cell_text(hdr2, "NO", bold=True, size=10)
    _shade(hdr2)
    _cell_text(nx.cell(6, 2), "NA", size=10, align="center")
    no_rows = [
        ("Percentage data capture (Hourly Values)", "{{ no.capture }}"),
        ("Hourly Maximum (ug/m³)", "{{ no.h_max }}"),
        ("Hourly Minimum (ug/m³)", "{{ no.h_min }}"),
        ("Daily average (ug/m³)", "{{ no.daily_avg }}"),
    ]
    for i, (a, b) in enumerate(no_rows):
        _cell_text(nx.cell(7 + i, 0), a, size=10)
        _cell_text(nx.cell(7 + i, 1), b, size=10, align="center")
    na1 = nx.cell(7, 2).merge(nx.cell(10, 2))
    _cell_text(na1, "NA", size=10, align="center")
    hdr3 = _merge_row(nx, 11, 0, 1)
    _cell_text(hdr3, "NOX", bold=True, size=10)
    _shade(hdr3)
    _cell_text(nx.cell(11, 2), "NA", size=10, align="center")
    nox_rows = [
        ("Percentage data capture (Hourly Values)", "{{ nox.capture }}"),
        ("Hourly Maximum (ug/m³)", "{{ nox.h_max }}"),
        ("Hourly Minimum (ug/m³)", "{{ nox.h_min }}"),
        ("Daily average (ug/m³)", "{{ nox.daily_avg }}"),
    ]
    for i, (a, b) in enumerate(nox_rows):
        _cell_text(nx.cell(12 + i, 0), a, size=10)
        _cell_text(nx.cell(12 + i, 1), b, size=10, align="center")
    na2 = nx.cell(12, 2).merge(nx.cell(15, 2))
    _cell_text(na2, "NA", size=10, align="center")
    _p(doc, "{{ nox_group.footnote }}", size=9, italic=True)
    _verdict_box(doc, "{{ nox_group.verdict_line }}")
    _figure(doc, "fig_no", "NO Hourly Concentration at the location.")
    _figure(doc, "fig_no2", "NO2 Hourly Concentration at the location.")
    _figure(doc, "fig_nox", "NOX Hourly Concentration at the location.")

    # 5.1.3 CO
    _heading(doc, "5.1.3 Carbon Monoxide (CO)", 3)
    _p(doc, "Carbon monoxide is a colorless and odorless gas which reduces the "
            "delivery of oxygen to the body's organs. For those with heart disease, "
            "exposure to low doses can result in chest pain. For healthier people, "
            "exposure to higher levels affects the central nervous system. "
            "Incomplete oxidation of fuel results in the formation of CO.",
       align="justify")
    _p(doc, "{{ co.narrative }}", align="justify")
    _caption(doc, "Table", "Summary of CO Results.")
    _summary_table(doc, [
        ("Percentage data capture (Hourly Values)", "{{ co.capture }}"),
        ("Hourly Maximum (ug/m³)", "{{ co.h_max }}"),
        ("Hourly Minimum (ug/m³)", "{{ co.h_min }}"),
        ("8 Hour Maximum (ug/m³)", "{{ co.r8_max }}"),
        ("8 Hour Minimum (ug/m³)", "{{ co.r8_min }}"),
        ("Hourly value > {{ co.limit_1h }} (ug/m³)", "{{ co.exceed_1h }}"),
        ("8-Hourly rolling average value > {{ co.limit_8h }} (ug/m³)",
         "{{ co.exceed_8h }}"),
        ("Daily average (ug/m³)", "{{ co.daily_avg }}"),
    ], [("1 Hour", "{{ co.limit_1h }}"), ("8 Hour", "{{ co.limit_8h }}")])
    _p(doc, "{{ co.footnote }}", size=9, italic=True)
    _verdict_box(doc, "{{ co.verdict_line }}")
    _figure(doc, "fig_co", "CO Hourly Concentration at the location.")
    _figure(doc, "fig_co8h", "CO 8 Hour Rolling Average Concentrations at the location.")

    # 5.1.4 H2S
    _heading(doc, "5.1.4 Hydrogen sulfide (H2S)", 3)
    _p(doc, "Hydrogen sulfide is a chemical compound with the formula H2S. It is a "
            "colorless chalcogen-hydride gas, and is poisonous, corrosive, and "
            "flammable, with trace amounts in ambient atmosphere having a "
            "characteristic foul odor of rotten eggs.", align="justify")
    _p(doc, "{{ h2s.narrative }}", align="justify")
    _caption(doc, "Table", "Summary of H₂S Results.")
    _summary_table(doc, [
        ("Percentage data capture (Hourly Values)", "{{ h2s.capture }}"),
        ("Hourly Maximum (µg/m³)", "{{ h2s.h_max }}"),
        ("Hourly Minimum (µg/m³)", "{{ h2s.h_min }}"),
        ("Hourly value > {{ h2s.limit_1h }} (ug/m³)", "{{ h2s.exceed_1h }}"),
        ("Daily Average (µg/m³)", "{{ h2s.daily_avg }}"),
    ], [("1 Hour", "{{ h2s.limit_1h }}"), ("24 Hour", "{{ h2s.limit_24h }}")])
    _p(doc, "{{ h2s.footnote }}", size=9, italic=True)
    _verdict_box(doc, "{{ h2s.verdict_line }}")
    _figure(doc, "fig_h2s", "H2S Hourly Concentration at the location.")

    # 5.1.5 O3
    _heading(doc, "5.1.5 Ozone (O3)", 3)
    _p(doc, "Ozone forms a protective layer which prevents entry of harmful "
            "ultraviolet radiation into the earth. The ground ozone is very harmful "
            "to human beings and the environment. It is released from industries, "
            "automobile emissions, gasoline vapors, solvents, chemicals, and "
            "electronic devices. Nitrogen oxides (NOx) and total Volatile Organic "
            "Compounds (TVOCs) also contribute to ground ozone formation. Ground "
            "ozone interferes with the plant's respiration process and enhances "
            "environmental stressor susceptibility. When ozone is inhaled by "
            "humans, reduced lung function, inflammation of airways, and irritation "
            "in the eyes, nose & throat are seen.", align="justify")
    _p(doc, "{{ o3.narrative }}", align="justify")
    _caption(doc, "Table", "Summary of O3 Result.")
    _summary_table(doc, [
        ("Percentage data capture (Hourly Values)", "{{ o3.capture }}"),
        ("Hourly Maximum (ug/m³)", "{{ o3.h_max }}"),
        ("8 Hour Maximum (ug/m³)", "{{ o3.r8_max }}"),
        ("8 Hour value > {{ o3.limit_8h }} (ug/m³)", "{{ o3.exceed_8h }}"),
        ("Daily average (ug/m³)", "{{ o3.daily_avg }}"),
    ], [("8 Hour", "{{ o3.limit_8h }}")])
    _p(doc, "{{ o3.footnote }}", size=9, italic=True)
    _verdict_box(doc, "{{ o3.verdict_line }}")
    _figure(doc, "fig_o3", "O3 Hourly Concentration at the location.")
    _figure(doc, "fig_o38h", "O3 8 Hour Rolling Average Concentrations at the location.")
    _figure(doc, "fig_no2_o3", "NO2 vs. O3 Hourly Concentrations at the location.")

    # 5.1.6 PM
    _heading(doc, "5.1.6 Particulate Matter (PM10 & PM2.5)", 3)
    _p(doc, "A mixture of particles with liquid droplets in the air forms "
            "particulate matter. PM10 are particles that have a size of less than "
            "or equal to 10 microns whereas PM2.5 are ultra-fine particles having a "
            "size of less than or equal to 2.5 microns. Particulate Matter is "
            "released from constructions, smoking, cleanings, renovations, "
            "demolitions, natural hazards such as earthquakes, volcanic eruptions, "
            "and emissions from industries such as brick kilns, paper & pulp, etc. "
            "These particles, when inhaled, can penetrate deeper into the "
            "respiratory system, and cause respiratory ailments such as asthma, "
            "coughing, sneezing, irritation in the airways, eyes, nose, throat "
            "irritation, etc. Studies have also shown links between PM exposure and "
            "diabetes.", align="justify")
    _p(doc, "{{ pm_group.narrative }}", align="justify")
    _caption(doc, "Table", "Summary of PM10 Results.")
    _summary_table(doc, [
        ("Percentage data capture (Hourly Values)", "{{ pm10.capture }}"),
        ("Hourly Maximum (ug/m³)", "{{ pm10.h_max }}"),
        ("Hourly Minimum (ug/m³)", "{{ pm10.h_min }}"),
        ("Daily Values > {{ pm10.limit_24h }} (ug/m³)", "{{ pm10.exceed_24h }}"),
        ("Daily average (ug/m³)", "{{ pm10.daily_avg }}"),
    ], [("24 Hour", "{{ pm10.limit_24h }}")])
    _p(doc, "{{ pm10.footnote }}", size=9, italic=True)
    _verdict_box(doc, "{{ pm10.verdict_line }}")
    _caption(doc, "Table", "Summary of PM2.5 Results.")
    _summary_table(doc, [
        ("Percentage data capture (Hourly Values)", "{{ pm25.capture }}"),
        ("Hourly Maximum (ug/m³)", "{{ pm25.h_max }}"),
        ("Hourly Minimum (ug/m³)", "{{ pm25.h_min }}"),
        ("Daily Values > {{ pm25.limit_24h }} (ug/m³)", "{{ pm25.exceed_24h }}"),
        ("Daily average (ug/m³)", "{{ pm25.daily_avg }}"),
    ], [("24 Hour", "{{ pm25.limit_24h }}")])
    _p(doc, "{{ pm25.footnote }}", size=9, italic=True)
    _verdict_box(doc, "{{ pm25.verdict_line }}")
    _figure(doc, "fig_pm10", "PM10 Hourly Concentrations at the location.")
    _figure(doc, "fig_pm25", "PM2.5 Hourly Concentrations at the location.")

    # Meteorology
    _p(doc, "- Meteorological Parameters Monitored result:", bold=True)
    _p(doc, "The evaluation and interpretation of gas emission measurements is only "
            "possible in comparison with meteorological data acquired concurrently. "
            "The structure of the atmosphere close to the ground is extremely "
            "important for the local climate. Knowing solar radiation as well as "
            "the air humidity and air temperature is necessary to evaluate chemical "
            "reactions of pollutants in the air.", align="justify")
    _heading(doc, "5.1.7 Temperature and humidity", 3)
    _p(doc, "Temperature and humidity play a significant role in gas emission "
            "measurements. The recorded data of temperature and humidity was "
            "captured for {{ monitoring_hours }} hours. The results for the location "
            "were summarized in the following table, and represented on a graph "
            "(Figures below).", align="justify")
    _heading(doc, "5.1.8 Barometric pressure", 3)
    _p(doc, "To predict the weather, it must be the first understanding how "
            "atmospheric pressure works. The higher the barometric pressure, the "
            "better it is for good weather conditions. Conversely, low pressures "
            "generally bring in more clouds and moisture, leading to poor "
            "visibility and even precipitation or snowfall. The recorded data of "
            "Barometric pressure was captured for {{ monitoring_hours }} hours.",
       align="justify")
    _heading(doc, "5.1.9 Wind speed and direction", 3)
    _p(doc, "Wind speed describes how fast the air is moving past a certain point. "
            "Wind direction describes the direction on a compass from which the "
            "wind emanates. Wind speed and direction are important for monitoring "
            "and predicting weather patterns and global climate. The recorded data "
            "of Wind speed and direction was captured for {{ monitoring_hours }} "
            "hours at the location. The results were summarized in the following "
            "table, and represented on a graph. Also, the Wind speed and direction "
            "were represented as a wind rose, and the wind frequency count and "
            "distribution was mentioned in the tables below. A wind rose is a "
            "graphic tool used by meteorologists to give a succinct view of how "
            "wind speed and direction are typically distributed at a particular "
            "location. Using a polar coordinate system of gridding, the frequency "
            "of winds over a period is plotted by wind direction, with color bands "
            "showing wind speed ranges. The direction of the longest spoke shows "
            "the wind direction with the greatest frequency.", align="justify")
    _caption(doc, "Table", "Monitored Meteorological Parameters result.")
    met = doc.add_table(rows=17, cols=2)
    met.style = "Table Grid"
    _polish(met)
    met_rows = [
        ("Ambient Temperature result", None),
        ("Percentage data capture (Hourly Values)", "{{ met.temp_capture }}"),
        ("Hourly Maximum (⁰C)", "{{ met.temp_max }}"),
        ("Hourly Minimum (⁰C)", "{{ met.temp_min }}"),
        ("Relative Humidity result", None),
        ("Percentage data capture (Hourly Values)", "{{ met.rh_capture }}"),
        ("Hourly Maximum (%)", "{{ met.rh_max }}"),
        ("Hourly Minimum (%)", "{{ met.rh_min }}"),
        ("Barometric Pressure result", None),
        ("Percentage data capture (Hourly Values)", "{{ met.pressure_capture }}"),
        ("Hourly Max (hPa)", "{{ met.pressure_max }}"),
        ("Hourly Minimum (hPa)", "{{ met.pressure_min }}"),
        ("Wind Parameters result", None),
        ("Percentage data capture (Hourly Values)", "{{ met.ws_capture }}"),
        ("Wind Speed Hourly Maximum (m/s)", "{{ met.ws_max }}"),
        ("Wind Speed Hourly Minimum (m/s)", "{{ met.ws_min }}"),
        ("Mean Wind Speed (m/s) / Prevailing Wind Direction",
         "{{ met.ws_mean }} / {{ met.prevailing }}"),
    ]
    for i, (a, b) in enumerate(met_rows):
        if b is None:
            merged = _merge_row(met, i, 0, 1)
            _cell_text(merged, a, bold=True, size=10)
            _shade(merged)
        else:
            _cell_text(met.cell(i, 0), a, size=10)
            _cell_text(met.cell(i, 1), b, size=10, align="center")
    for key, cap_text in [
        ("fig_temp", "Hourly Temperature at the location."),
        ("fig_rh", "Hourly Relative Humidity at the location."),
        ("fig_pressure", "Hourly Pressure at the location."),
        ("fig_ws", "Hourly Wind Speed at the location."),
        ("fig_windrose", "Wind Rose at the location."),
        ("fig_windclassfreq", "Wind class frequency distribution graph at the location."),
    ]:
        _figure(doc, key, cap_text)

    # Wind rose over the satellite tile. Wrapped in a condition because the
    # site map can be absent — no coordinates, no API key, a failed fetch —
    # and an empty figure with a live caption would misnumber everything
    # after it in the List of Figures.
    _p(doc, "{%p if fig_windrose_map %}", size=1, space_after=0)
    _p(doc, "{{ fig_windrose_map }}", align="center")
    _caption(doc, "Figure", "Wind Rose over the monitoring location.")
    _p(doc, "{%p endif %}", size=1, space_after=0)

    # Wind tables 14/15 — dynamic columns via dedicated {%tc %} cells
    for cap_text, rows_key, totals_key in [
        ("Wind class frequency distribution at the location.", "wind_pct_rows",
         "wind_pct_totals"),
        ("Wind class count at the location.", "wind_count_rows", "wind_count_totals"),
    ]:
        _caption(doc, "Table", cap_text)
        wt = doc.add_table(rows=7, cols=5)
        wt.style = "Table Grid"
        _polish(wt)
        # header row: label | tc-for | {{c}} | tc-endfor | Total
        _cell_text(wt.cell(0, 0), "Directions / Wind Classes (m/s)", bold=True, size=9)
        _shade(wt.cell(0, 0))
        wt.cell(0, 1).paragraphs[0].text = "{%tc for c in wind_class_labels %}"
        _cell_text(wt.cell(0, 2), "{{ c }}", bold=True, size=9, align="center")
        _shade(wt.cell(0, 2))
        wt.cell(0, 3).paragraphs[0].text = "{%tc endfor %}"
        _cell_text(wt.cell(0, 4), "Total", bold=True, size=9, align="center")
        _shade(wt.cell(0, 4))
        # looped data rows
        _tr_tag_row(wt, 1, "{%%tr for r in %s %%}" % rows_key)
        _cell_text(wt.cell(2, 0), "{{ r.direction }}", size=9)
        wt.cell(2, 1).paragraphs[0].text = "{%tc for v in r.vals %}"
        _cell_text(wt.cell(2, 2), "{{ v }}", size=9, align="center")
        wt.cell(2, 3).paragraphs[0].text = "{%tc endfor %}"
        _cell_text(wt.cell(2, 4), "{{ r.total }}", size=9, align="center")
        _tr_tag_row(wt, 3, "{%tr endfor %}")
        # sub-total row
        _cell_text(wt.cell(4, 0), "Sub-Total", bold=True, size=9)
        wt.cell(4, 1).paragraphs[0].text = "{%%tc for v in %s %%}" % totals_key
        _cell_text(wt.cell(4, 2), "{{ v }}", size=9, align="center")
        wt.cell(4, 3).paragraphs[0].text = "{%tc endfor %}"
        _cell_text(wt.cell(4, 4), "{{ %s_grand }}" % totals_key, size=9,
                   align="center")
        # calms / missing rows (span value columns)
        _cell_text(wt.cell(5, 0), "Calms", size=9)
        c5 = wt.cell(5, 1).merge(wt.cell(5, 4))
        _cell_text(c5, "{{ %s_calms }}" % totals_key, size=9, align="center")
        _cell_text(wt.cell(6, 0), "Missing/Incomplete", size=9)
        c6 = wt.cell(6, 1).merge(wt.cell(6, 4))
        _cell_text(c6, "{{ %s_missing }}" % totals_key, size=9, align="center")
        _p(doc, space_after=6)

    # --- 6. Conclusions
    _heading(doc, "6. Conclusions", 1)
    _p(doc, "Key observations arising from the examination of the recorded data for "
            "the monitoring period in the project site "
            "({{ monitoring_window_text }}).", align="justify")
    doc.add_paragraph("The average data capture for the station was "
                      "{{ overall_capture }} % for air quality and meteorological "
                      "parameters.", style="List Bullet")
    p = doc.add_paragraph(style="Normal")
    p.add_run("{%p for c in conclusion_blocks %}")
    doc.add_paragraph("For {{ c.title }} concentrations:", style="List Bullet")
    p2 = doc.add_paragraph(style="Normal")
    p2.add_run("{%p for line in c.lines %}")
    doc.add_paragraph("- {{ line }}", style="List Bullet 2")
    p3 = doc.add_paragraph(style="Normal")
    p3.add_run("{%p endfor %}")
    p4 = doc.add_paragraph(style="Normal")
    p4.add_run("{%p endfor %}")
    doc.add_paragraph("{{ met_conclusion_1 }}", style="List Bullet")
    doc.add_paragraph("{{ met_conclusion_2 }}", style="List Bullet")
    doc.add_paragraph("{{ met_conclusion_3 }}", style="List Bullet")
    doc.add_paragraph("{{ met_conclusion_4 }}", style="List Bullet")
    _p(doc, "The prevailing wind direction at the site was {{ met.prevailing }}.")

    # --- 7. Recommendations
    _heading(doc, "7. Recommendations", 1)
    _p(doc, "The following recommendations arise from the results of this "
            "survey.", align="justify", space_after=8)
    p = doc.add_paragraph(style="Normal")
    p.add_run("{%p for r in recommendations %}")
    doc.add_paragraph("{{ r }}", style="List Bullet")
    p = doc.add_paragraph(style="Normal")
    p.add_run("{%p endfor %}")

    # --- 8. Measurement uncertainty
    _heading(doc, "8. Measurement Uncertainty", 1)
    _p(doc, "{{ uncertainty_text }}", align="justify")

    # --- 9. Limitations and reliance
    _heading(doc, "9. Limitations", 1)
    p = doc.add_paragraph(style="Normal")
    p.add_run("{%p for l in limitations %}")
    doc.add_paragraph("{{ l }}", style="List Bullet")
    p = doc.add_paragraph(style="Normal")
    p.add_run("{%p endfor %}")

    # --- Appendices
    doc.add_page_break()
    _heading(doc, "Appendix 1 Valid Data Exception", 1)
    _p(doc, "{{ appendix1_text }}", align="justify")
    _heading(doc, "Appendix 2 Valid Data Terms", 1)
    for lead, body in [
        ("Span / zero check.", " A manual zero calibration check is performed "
         "whereby air is passed through filter element, removing particulates, "
         "before entering the sensor in the analyzer. Data is invalidated when "
         "these checks occur."),
        ("Multipoint Calibration.", " To perform multipoint calibration, span and "
         "zero gases are passed through filter element, removing particulates, "
         "before entering the sensor in the analyzer. Data is invalidated when "
         "calibration occurs."),
        ("Instrument fault", " refers to a period when the instrument was not in "
         "the normal operating mode and did not measure a representative value of "
         "the existing conditions."),
        ("Data Communication Issue", " refers to a period when instrument is not "
         "connected to data logger (Configuration lost)."),
        ("Power Interruption", " refers to no power to the AAQMS therefore no data "
         "was collected at that time."),
    ]:
        pp = _p(doc, "", align="justify")
        rr = pp.add_run(lead)
        rr.bold = True
        pp.add_run(body)
    _heading(doc, "Appendix 3 Calibration certificates", 1)
    _p(doc, "{%p if cert_rows %}", size=1, space_after=0)
    _p(doc, "The gaseous analysers were calibrated by an accredited "
            "calibration laboratory against internationally traceable "
            "reference standards prior to the survey. The certificates are "
            "summarised below and reproduced in full on the following pages.",
       align="justify")
    _caption(doc, "Table", "Calibration Certificate Summary")
    ct = doc.add_table(rows=4, cols=6)
    ct.style = "Table Grid"
    _polish(ct)
    for j, h in enumerate(["CERTIFICATE No.", "PARAMETER", "MODEL / S.N.",
                           "CALIBRATION DATE", "DUE DATE", "RESULT"]):
        _cell_text(ct.cell(0, j), h, bold=True, size=8.5, align="center")
        _shade(ct.cell(0, j))
    _tr_tag_row(ct, 1, "{%tr for c in cert_rows %}")
    crow = ct.rows[2]
    for k, key in enumerate(["number", "parameter", "model_sn", "date",
                             "due_date"]):
        _cell_text(crow.cells[k], "{{ c.%s }}" % key, size=9,
                   align="center" if k else "left")
    _cell_text(crow.cells[5], "{{ c.result }}", size=9, align="center",
               bold=True)
    _tr_tag_row(ct, 3, "{%tr endfor %}")
    _p(doc, space_after=8)
    _p(doc, "{%p endif %}", size=1, space_after=0)
    _p(doc, "{%p if calibration_images %}", size=1, space_after=0)
    _p(doc, "{%p for c in calibration_images %}", size=1, space_after=0)
    _p(doc, "{{ c.title }}", bold=True, size=10, color=NAVY, space_after=4)
    _p(doc, "{{ c.image }}", align="center")
    _p(doc, "{%p endfor %}", size=1, space_after=0)
    _p(doc, "{%p else %}", size=1, space_after=0)
    _p(doc, "[Calibration certificates to be attached — upload scanned "
            "certificates for this campaign.]", italic=True)
    _p(doc, "{%p endif %}", size=1, space_after=0)
    _heading(doc, "Appendix 4 Environmental license for the institution", 1)
    _p(doc, "{%p if license_images %}", size=1, space_after=0)
    _p(doc, "{%p for img in license_images %}", size=1, space_after=0)
    _p(doc, "{{ img }}", align="center")
    _p(doc, "{%p endfor %}", size=1, space_after=0)
    _p(doc, "{%p else %}", size=1, space_after=0)
    _p(doc, "[Environmental license to be attached — upload scanned license for "
            "this provider.]", italic=True)
    _p(doc, "{%p endif %}", size=1, space_after=0)

    # End-of-report marker. A reader who receives a loose bundle needs to
    # know they have the whole document; without it a report that ends on a
    # scanned licence page is indistinguishable from one that is missing
    # pages.
    #
    # One marker only. A second, unbordered version of this block was added
    # here at some point and the original was never removed, so every report
    # closed with the words twice, six millimetres apart.
    _p(doc, "", size=8)
    end = doc.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    end.paragraph_format.space_before = Pt(10)
    ePr = end._p.get_or_add_pPr()
    eBdr = OxmlElement("w:pBdr")
    for side in ("top", "bottom"):
        e = OxmlElement("w:%s" % side)
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "6")
        e.set(qn("w:color"), BLUE_FILL)
        e.set(qn("w:space"), "6")
        eBdr.append(e)
    ePr.append(eBdr)
    er = end.add_run("END OF REPORT")
    er.bold = True
    er.font.size = Pt(10)
    er.font.color.rgb = NAVY
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("{{ report_number }}  \u00b7  Rev {{ revision }}  \u00b7  "
                     "{{ project_name }}")
    sr.font.size = Pt(8)
    sr.font.color.rgb = MUTED_GREY

    stats = _typeset(doc)
    _update_fields_on_open(doc)
    _modernise_settings(doc)
    doc.save(out_path)
    import logging
    logging.getLogger(__name__).info(
        "typesetting: %s", ", ".join(f"{k}={v}" for k, v in stats.items()))
    return out_path


if __name__ == "__main__":
    path = build()
    print(f"Template written: {path}")
