# -*- coding: utf-8 -*-
"""Build the noise monitoring report DOCX.

Unlike the air report, which fills a docxtpl master template, this generator
writes the document directly with python-docx. A noise report is a fifth the
size of an air report and direct construction keeps every design decision in
one readable file.

It still plugs into the shared machinery:

* Headings use the built-in "Heading 1/2" styles and captions the "Caption"
  style, and three TOC anchor fields are written up front — exactly what
  ``report/fields.py`` scans for, so the same bookmark/PAGEREF indexing that
  fixed the air report's page numbers builds this one's Table of Contents,
  List of Figures and List of Tables.
* The watermark is applied through ``report.template_builder._watermark``,
  the same VML shape as the air report.
* Charts come from ``report/noise_charts.py`` and imaging, storage,
  versioning, review and the on-screen reader all operate on the produced
  file exactly as they do for air.
"""
from __future__ import annotations

import logging
import os
from datetime import datetime
from typing import Dict, List, Optional

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

from noise_calc import (CONSTRUCTION_CORRECTIONS, NOISE_LIMITS, NoiseSummary,
                        NoiseVerdict)

log = logging.getLogger(__name__)

NAVY = RGBColor(0x0F, 0x3D, 0x6E)
NAVY_HEX = "0F3D6E"
GREEN_HEX = "1E7D4F"
AMBER_HEX = "B06A00"
MUT = RGBColor(0x6B, 0x72, 0x80)
INK = RGBColor(0x20, 0x24, 0x2B)
HAIR = "D9DDE3"
ROW = "F4F6F9"

ASSETS = os.path.join(os.path.dirname(os.path.abspath(__file__)), "assets")


# ---------------------------------------------------------------------------
# small helpers
# ---------------------------------------------------------------------------
def _bleed(paragraph):
    """Turn the paragraph's inline picture into a page-anchored one at 0,0.

    An inline picture at full page height is laid out inside a line box, and
    the line box aligns it to the baseline: the top is clipped and white is
    left at the foot. Forcing the leading to the page depth does not fix it,
    it only moves the clipping. Anchoring the picture to the page, behind the
    text, takes it out of the line box altogether, which is how a full-bleed
    cover is supposed to be built and how both Word and LibreOffice place it
    identically.
    """
    inline = paragraph._p.find(
        ".//" + qn("w:drawing") + "/" + qn("wp:inline"))
    if inline is None:
        return
    anchor = OxmlElement("wp:anchor")
    for k, v in (("distT", "0"), ("distB", "0"), ("distL", "0"),
                 ("distR", "0"), ("simplePos", "0"), ("relativeHeight", "1"),
                 ("behindDoc", "1"), ("locked", "0"), ("layoutInCell", "1"),
                 ("allowOverlap", "1")):
        anchor.set(k, v)
    sp = OxmlElement("wp:simplePos")
    sp.set("x", "0")
    sp.set("y", "0")
    anchor.append(sp)
    for tag, rel in (("wp:positionH", "page"), ("wp:positionV", "page")):
        pos = OxmlElement(tag)
        pos.set("relativeFrom", rel)
        off = OxmlElement("wp:posOffset")
        off.text = "0"
        pos.append(off)
        anchor.append(pos)
    for child in list(inline):
        tag = child.tag.split("}")[-1]
        if tag == "extent":
            anchor.append(child)
            eff = OxmlElement("wp:effectExtent")
            for e in ("l", "t", "r", "b"):
                eff.set(e, "0")
            anchor.append(eff)
            anchor.append(OxmlElement("wp:wrapNone"))
        elif tag != "effectExtent":
            anchor.append(child)
    drawing = inline.getparent()
    drawing.remove(inline)
    drawing.append(anchor)


def _tight(p, before=0, after=6):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    return p


def _table_width(table, mm: float):
    """Force the table's overall width. Cell widths alone leave the table
    itself at its default, which is why the cover's footer band stopped
    short of the paper edge instead of bleeding across it."""
    tblPr = table._tbl.tblPr
    for old in tblPr.findall(qn("w:tblW")):
        tblPr.remove(old)
    w = OxmlElement("w:tblW")
    w.set(qn("w:w"), str(int(mm * 56.7)))      # twentieths of a point
    w.set(qn("w:type"), "dxa")
    tblPr.append(w)


def _shade(cell, hexval):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexval)
    tcPr.append(sh)


def _borders(cell, **spec):
    tcPr = cell._tc.get_or_add_tcPr()
    tb = OxmlElement("w:tcBorders")
    for tag in ("top", "bottom", "left", "right"):
        el = OxmlElement("w:" + tag)
        v = spec.get(tag)
        if v is None:
            el.set(qn("w:val"), "nil")
        else:
            colour, sz = v
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), str(sz))
            el.set(qn("w:color"), colour)
        tb.append(el)
    tcPr.append(tb)


def _cell(cell, text, size=11, bold=False, colour=INK, align="left"):
    p = cell.paragraphs[0]
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT,
                   "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    _tight(p, 2, 2)
    r = p.add_run(str(text))
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.name = "Times New Roman"
    r.font.color.rgb = colour
    return p


def _table(doc, headers, rows, widths_mm, aligns=None, size=None):
    """Tables at the company's 11 pt, stepped down when a table is wide.

    Their reports carry tables of two to four columns at 11-12 pt. This
    report has a nine-column summary of measured levels, and 11 pt in a
    nine-column table wraps every heading onto three lines. The size follows
    the column count so the narrow tables match theirs and the wide ones stay
    readable instead of matching a table that does not exist in their set.
    """
    if size is None:
        n = len(headers)
        size = 9.5 if n >= 8 else (10.5 if n >= 6 else 11)
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for j, wmm in enumerate(widths_mm):
        for c in t.columns[j].cells:
            c.width = Mm(wmm)
    for j, h in enumerate(headers):
        c = t.cell(0, j)
        _shade(c, NAVY_HEX)
        _cell(c, h, size=size, bold=True, colour=RGBColor(255, 255, 255),
              align="center")
        _borders(c, top=(NAVY_HEX, 4), bottom=(NAVY_HEX, 4))
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            c = t.cell(i + 1, j)
            if i % 2 == 1:
                _shade(c, ROW)
            _cell(c, v, size=size,
                  align=(aligns[j] if aligns else "center"))
            _borders(c, bottom=(HAIR, 2))
    return t


def _callout(doc, text, colour_hex, fill_hex, mark):
    t = doc.add_table(rows=1, cols=1)
    t.autofit = False
    body = t.cell(0, 0)
    body.width = Mm(164)
    _shade(body, fill_hex)
    tcPr = body._tc.get_or_add_tcPr()
    tb = OxmlElement("w:tcBorders")
    for tag in ("top", "bottom", "right"):
        el = OxmlElement("w:" + tag)
        el.set(qn("w:val"), "nil")
        tb.append(el)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:color"), colour_hex)
    tb.append(left)
    tcPr.append(tb)
    p = _tight(body.paragraphs[0], 3, 3)
    r = p.add_run(f"{mark}  ")
    r.font.size = Pt(9)
    r.font.bold = True
    r.font.color.rgb = RGBColor(int(colour_hex[0:2], 16),
                                int(colour_hex[2:4], 16),
                                int(colour_hex[4:6], 16))
    r2 = p.add_run(text)
    r2.font.size = Pt(9)
    _tight(doc.add_paragraph(), 0, 2)


def _heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    if level == 1:
        pPr = h._p.get_or_add_pPr()
        b = OxmlElement("w:pBdr")
        bt = OxmlElement("w:bottom")
        bt.set(qn("w:val"), "single")
        bt.set(qn("w:sz"), "8")
        bt.set(qn("w:color"), NAVY_HEX)
        b.append(bt)
        pPr.append(b)
    return h


def _caption(doc, text):
    p = doc.add_paragraph(text, style="Caption")
    _tight(p, 4, 8)
    for r in p.runs:
        # Not bold: theirs are not, and at 13 pt a bold caption reads as a
        # heading and competes with the section titles around it.
        r.font.bold = False
        r.font.size = Pt(13)
        r.font.color.rgb = NAVY
    return p


def _muted(doc, text, size=10):
    p = _tight(doc.add_paragraph(), 0, 8)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.italic = True
    r.font.color.rgb = MUT


def _bullets(doc, items):
    """A numbered procedure an auditor can tick off, rather than prose —
    the form BSA's manual reports use for method statements."""
    for it in items:
        p = doc.add_paragraph(it, style="List Bullet")
        _tight(p, 0, 2)
        for r in p.runs:
            r.font.size = Pt(13)
            r.font.name = "Times New Roman"
    return doc


def _body(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def _toc_anchor(doc, instr):
    """An empty TOC field: begin / instruction / separate / (blank) / end.
    ``report.fields.build_indexes`` finds these and replaces each with real
    entries carrying bookmarked page numbers."""
    p = _tight(doc.add_paragraph(), 0, 0)
    for kind, txt in (("begin", None), (None, instr), ("separate", None),
                      (None, None), ("end", None)):
        r = OxmlElement("w:r")
        if kind:
            f = OxmlElement("w:fldChar")
            f.set(qn("w:fldCharType"), kind)
            r.append(f)
        elif txt:
            it = OxmlElement("w:instrText")
            it.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            it.text = txt
            r.append(it)
        else:
            t = OxmlElement("w:t")
            t.text = ""
            r.append(t)
        p._p.append(r)
    return p


def _fmt_dt(dt: Optional[datetime]) -> str:
    if not dt:
        return "—"
    return dt.strftime("%B %-d, %Y, %I:%M %p") if os.name != "nt" \
        else dt.strftime("%B %d, %Y, %I:%M %p")


def _f(v) -> str:
    return "—" if v is None else f"{v:.2f}"


def _interval_phrase(seconds: float) -> str:
    """Describe the logging interval the meter actually used.

    This was hardcoded as "one-minute" throughout the report and became a
    false statement about the measurement method the first time a
    per-second logger was used — in a document that goes to a regulator.
    """
    if seconds <= 1.5:
        return "one-second"
    if seconds < 55:
        return f"{int(round(seconds))}-second"
    if seconds <= 65:
        return "one-minute"
    mins = seconds / 60.0
    return (f"{int(round(mins))}-minute" if abs(mins - round(mins)) < 0.05
            else f"{mins:.1f}-minute")


def _date_only(dt: Optional[datetime]) -> str:
    """The date with its year. Splitting the long form on its first comma
    dropped the year and printed "July 14" on the cover."""
    if not dt:
        return "—"
    return (dt.strftime("%B %-d, %Y") if os.name != "nt"
            else dt.strftime("%B %d, %Y"))


# ---------------------------------------------------------------------------
# the document
# ---------------------------------------------------------------------------
def generate_noise_report(campaign, summary: NoiseSummary,
                          verdicts: List[NoiseVerdict],
                          figs: Dict[str, str], out_path: str,
                          site_map_path: Optional[str] = None,
                          site_photo_paths: Optional[List[str]] = None,
                          cover_photo_path: Optional[str] = None,
                          equipment_photo_paths: Optional[List[str]] = None,
                          calibration_items: Optional[List[dict]] = None,
                          license_image_paths: Optional[List[str]] = None,
                          work_dir: Optional[str] = None) -> str:
    doc = Document()
    st = doc.styles["Normal"]
    # Sizes measured from BSA's own report (BR-R-220726-104) rather than
    # chosen: body 13 pt, Heading 1 16 pt, Heading 2 14 pt, Heading 3 13 pt,
    # captions 13 pt, single leading. The generated report previously ran at
    # 11 pt on 1.15 leading, which is why it read as cramped beside theirs.
    st.font.name = "Times New Roman"
    st.font.size = Pt(13)
    st.paragraph_format.space_after = Pt(6)
    st.paragraph_format.line_spacing = 1.0
    for lvl, sz in (("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 13)):
        h = doc.styles[lvl]
        h.font.name = "Times New Roman"
        h.font.size = Pt(sz)
        h.font.bold = True
        h.font.color.rgb = NAVY
    cap = doc.styles["Caption"]
    cap.font.name = "Times New Roman"
    cap.font.color.rgb = NAVY

    # Table numbers are counted rather than written in, because optional
    # content — meteorology here, a site map or photographs elsewhere —
    # changes how many tables come before any given one.
    _tn = {"n": 0}

    def tnum() -> int:
        _tn["n"] += 1
        return _tn["n"]

    # Appendix letters are settled before the body is written, because the
    # methodology refers to the certificate appendix by letter long before
    # that appendix is reached. Fixed letters left gaps — a report with no
    # certificates began at Appendix C — and a letter worked out at the end
    # cannot be quoted at the start.
    _appendix: Dict[str, str] = {}
    _next = 0
    if calibration_items:
        _appendix["certificates"] = chr(ord("A") + _next)
        _next += 1
    if license_image_paths:
        _appendix["licence"] = chr(ord("A") + _next)
        _next += 1

    _fg = {"n": 0}

    def _fn() -> int:
        """Figures counted in document order. Working the number out from
        which optional blocks exist was fragile the moment another optional
        figure — the equipment photograph — appeared before the charts."""
        _fg["n"] += 1
        return _fg["n"]

    from report.imaging import slim
    wd = work_dir or os.path.dirname(os.path.abspath(out_path))

    def pic(p, path, width_mm):
        p.add_run().add_picture(slim(path, wd) or path, width=Mm(width_mm))

    # ---- cover -----------------------------------------------------------
    # One full-bleed picture. Every field on it is repeated on the document
    # control page, which is selectable text, so nothing is lost to search.
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Mm(210), Mm(297)
    sec.top_margin = sec.bottom_margin = Mm(0)
    sec.left_margin = sec.right_margin = Mm(0)
    sec.header_distance = sec.footer_distance = Mm(0)
    sec.gutter = Mm(0)

    logo = os.path.join(ASSETS, "logo_left.png")

    plate_done = False
    try:
        from report.noise_cover import date_range, render_cover
        plate = os.path.join(wd, "noise_cover_plate.png")
        render_cover(
            plate,
            survey_dates=[date_range(campaign.monitoring_start,
                                     campaign.monitoring_end),
                          date_range(campaign.monitoring_start,
                                     campaign.monitoring_end, abbrev=True)],
            location=(campaign.site_name or campaign.project_name or "—"),
            client=campaign.client or "—",
            report_number=campaign.report_number or "—",
            revision=campaign.revision or "00",
            issue_date=_date_only(campaign.reporting_date),
            photo_path=cover_photo_path,
            logo_path=logo if os.path.exists(logo) else None)
        p = _tight(doc.add_paragraph(), 0, 0)
        p.add_run().add_picture(plate, width=Mm(210), height=Mm(297))
        _bleed(p)
        for r in p.runs:
            r.font.size = Pt(1)
        plate_done = True
    except Exception:  # noqa: BLE001
        log.warning("cover plate failed — falling back to a plain cover",
                    exc_info=True)

    if not plate_done:
        # A plain cover carrying the same facts, so a rendering failure costs
        # the design and not the document.
        for _ in range(6):
            _tight(doc.add_paragraph(), 0, 0)
        for text, size, bold, colour in (
                ("NOISE MONITORING REPORT", 26, True, NAVY),
                (campaign.project_name or "", 15, True, NAVY),
                (f"{campaign.site_name or ''}  ·  24-hour attended noise "
                 f"survey", 10, False, MUT)):
            p = _tight(doc.add_paragraph(), 0, 6)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text)
            r.font.size = Pt(size)
            r.font.bold = bold
            r.font.color.rgb = colour
        info = [("Client", campaign.client or "—"),
                ("Monitoring location", campaign.site_name or "—"),
                ("Survey period",
                 f"{_fmt_dt(campaign.monitoring_start)} to "
                 f"{_fmt_dt(campaign.monitoring_end)}"),
                ("Report number", campaign.report_number or "—"),
                ("Revision and issue date",
                 f"{campaign.revision or '00'}    ·    "
                 f"{_date_only(campaign.reporting_date)}")]
        t = doc.add_table(rows=len(info), cols=2)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.autofit = False
        _table_width(t, 180)
        for i, (lab, val) in enumerate(info):
            a, b = t.cell(i, 0), t.cell(i, 1)
            a.width = Mm(52)
            b.width = Mm(128)
            _cell(a, lab, size=8.5, colour=MUT)
            _cell(b, val, size=10, bold=True)
            for c in (a, b):
                _borders(c, bottom=(HAIR, 2))

    # ---- body section ----------------------------------------------------
    body = doc.add_section(WD_SECTION.NEW_PAGE)
    body.top_margin = body.bottom_margin = Mm(22)
    body.left_margin = body.right_margin = Mm(23)

    hdr = body.header
    hdr.is_linked_to_previous = False
    ht = hdr.add_table(rows=1, cols=2, width=Mm(164))
    ht.autofit = False
    hc1, hc2 = ht.cell(0, 0), ht.cell(0, 1)
    hc1.width = Mm(50)
    hc2.width = Mm(114)
    hp = _tight(hc1.paragraphs[0], 0, 0)
    if os.path.exists(logo):
        hp.add_run().add_picture(logo, width=Mm(22))
    _cell(hc2, "Noise Monitoring Report", size=10, bold=True, colour=NAVY,
          align="right")
    p2 = _tight(hc2.add_paragraph(), 0, 0)
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p2.add_run("  ·  ".join(
        x for x in (campaign.project_name, campaign.report_number,
                    f"Rev {campaign.revision or '00'}") if x))
    r.font.size = Pt(7.5)
    r.font.color.rgb = MUT
    for c in (hc1, hc2):
        _borders(c, bottom=(NAVY_HEX, 6))

    ftr = body.footer
    ftr.is_linked_to_previous = False
    fp = _tight(ftr.paragraphs[0], 0, 0)
    r = fp.add_run("CONFIDENTIAL · Bander Said Allehiany for Environmental "
                   "Consultancy")
    r.font.size = Pt(7)
    r.font.color.rgb = MUT
    fp2 = _tight(ftr.add_paragraph(), 0, 0)
    fp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = fp2.add_run("Page ")
    rr.font.size = Pt(8)
    rr.font.color.rgb = NAVY
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), " PAGE ")
    fp2._p.append(fld)

    try:
        from report.template_builder import _watermark
        wm = os.path.join(ASSETS, "watermark.png")
        if os.path.exists(wm):
            _watermark(body, wm)
    except Exception:  # noqa: BLE001
        log.warning("watermark skipped", exc_info=True)

    # ---- document control ------------------------------------------------
    p = _tight(doc.add_paragraph(), 12, 2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Noise Monitoring Report")
    r.font.size = Pt(17)
    r.font.bold = True
    p = _tight(doc.add_paragraph(), 0, 2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("for").font.size = Pt(11)
    p = _tight(doc.add_paragraph(), 0, 8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(campaign.project_name or "")
    r.font.size = Pt(14)
    r.font.bold = True

    _table(doc, ["Project", campaign.project_name or "—"],
           [["Document Title", "Noise Monitoring Report"],
            ["Client", campaign.client or "—"],
            ["Report Number", campaign.report_number or "—"]],
           [52, 112], aligns=["left", "left"])

    _caption(doc, "Revision History")
    _table(doc, ["Rev", "Reporting Date", "Prepared By",
                 "Project Supervision", "Status"],
           [[campaign.revision or "00",
             _date_only(campaign.reporting_date),
             campaign.prepared_by or "—",
             campaign.project_supervision or "—",
             campaign.document_status or "Issued for Client Use"]],
           [18, 36, 38, 38, 34])

    # ---- indexes ---------------------------------------------------------
    doc.add_page_break()
    _heading(doc, "Table of Contents", 1)
    _toc_anchor(doc, ' TOC \\o "1-3" \\h \\z \\u ')
    _heading(doc, "List of Figures", 1)
    _toc_anchor(doc, ' TOC \\h \\z \\c "Figure" ')
    _heading(doc, "List of Tables", 1)
    _toc_anchor(doc, ' TOC \\h \\z \\c "Table" ')

    # ---- executive summary ----------------------------------------------
    doc.add_page_break()
    _heading(doc, "Executive Summary", 1)
    day = f"{summary.day_start_hour:02d}:00–{summary.day_end_hour:02d}:00"
    interval = _interval_phrase(getattr(summary, "interval_seconds", 60.0))
    _body(doc,
          f"Bander Said Allehiany (BSA) was commissioned by "
          f"{campaign.client} to conduct an attended environmental noise "
          f"survey at {campaign.project_name}. Sound pressure levels were "
          f"logged continuously at {interval} intervals at one location, "
          f"from {_fmt_dt(campaign.monitoring_start)} to "
          f"{_fmt_dt(campaign.monitoring_end)}. The location was selected "
          f"to represent background conditions at the project site; key "
          f"observations such as unusual sounds and passing plant were "
          f"recorded by the attending team throughout the survey.")

    stats = [("LAeq, T", summary.laeq_t), ("L Day", summary.l_day),
             ("L Night", summary.l_night), ("LA10", summary.la10),
             ("LA50", summary.la50), ("LA90", summary.la90),
             ("Lmax", summary.lmax), ("Lmin", summary.lmin)]
    t = doc.add_table(rows=2, cols=8)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for j, (lab, val) in enumerate(stats):
        a, b = t.cell(0, j), t.cell(1, j)
        for c in (a, b):
            c.width = Mm(20.5)
        _cell(a, lab, size=7, colour=MUT, align="center")
        _cell(b, "—" if val is None else f"{val:.1f}", size=13, bold=True,
              colour=NAVY, align="center")
        _borders(a, top=(HAIR, 4), left=(HAIR, 4), right=(HAIR, 4))
        _borders(b, bottom=(HAIR, 4), left=(HAIR, 4), right=(HAIR, 4))
    _muted(doc,
           f"All values in dB(A) · derived from "
           f"{summary.valid_records:,} validated {interval} intervals · "
           f"day period {day} · data capture {summary.data_capture_pct}%")

    if verdicts:
        worst = max(verdicts, key=lambda v: v.margin)
        if worst.status == "ok":
            _callout(doc, "The measured day-time and night-time levels "
                     "complied with the applicable NCEC limits throughout "
                     "the survey period.", GREEN_HEX, "EEF6F1",
                     "KEY FINDING")
        else:
            _callout(doc, "One or more measured levels exceeded the "
                     "applicable NCEC limit; the assessment in Section 6 "
                     "sets out the margins and the applicable corrections.",
                     AMBER_HEX, "FDF3E6", "KEY FINDING")
    else:
        _callout(doc, "The measured levels are stated against every NCEC "
                 "land-use category in Section 6; the applicable category "
                 "for this location is to be determined by the consultant.",
                 NAVY_HEX, "EEF3F9", "KEY FINDING")

    # ---- introduction and scope -----------------------------------------
    _heading(doc, "1  Introduction", 1)
    _body(doc,
          f"A noise monitoring survey was conducted at "
          f"{campaign.project_name} to develop 24-hour noise levels at one "
          f"location at the project site. The stored data was used to "
          f"derive the key sound level parameters LAeq, LA10, LA50, LA90, "
          f"L Day and L Night, and the results are compared with the "
          f"standards of the National Center for Environmental Compliance "
          f"(NCEC). Graphical representations of the monitored record are "
          f"provided in Section 6.")

    _heading(doc, "2  Scope of Work", 1)
    _body(doc,
          f"A noise baseline survey using a Class 1 sound level meter was "
          f"conducted to characterise the noise environment. Monitoring was "
          f"carried out continuously, with instantaneous sound pressure "
          f"levels sampled and stored at {interval} intervals. The "
          f"continuous equivalent level (LAeq) over the day and night "
          f"periods was derived from the logged record by energy averaging.")
    _caption(doc, f"Table {tnum()} — Time of monitoring")
    dur = (campaign.monitoring_end - campaign.monitoring_start)
    _table(doc, ["Site ID", "Start", "End", "Duration (hrs)"],
           [["N1", _fmt_dt(campaign.monitoring_start),
             _fmt_dt(campaign.monitoring_end),
             f"{dur.total_seconds() / 3600:.0f}"]],
           [22, 56, 56, 30])
    _caption(doc, f"Table {tnum()} — Coordinates of the noise monitoring location")
    _table(doc, ["Site", "Geographical Coordinates"],
           [["N1", f"N {campaign.latitude:.6f}   "
                   f"E {campaign.longitude:.6f}"]],
           [40, 124])

    if site_map_path and os.path.exists(site_map_path):
        p = _tight(doc.add_paragraph(), 6, 0)
        pic(p, site_map_path, 150)
        _caption(doc, f"Figure {_fn()} — Noise monitoring location")

    if site_photo_paths:
        imgs = [p for p in site_photo_paths if p and os.path.exists(p)]
        if imgs:
            t = doc.add_table(rows=(len(imgs) + 1) // 2, cols=2)
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            for k, path in enumerate(imgs[:4]):
                cell = t.cell(k // 2, k % 2)
                cp = _tight(cell.paragraphs[0], 2, 2)
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pic(cp, path, 74)
            _caption(doc, f"Figure {_fn()} — Noise monitoring at the site")

    # ---- methodology -----------------------------------------------------
    _heading(doc, "3  Monitoring Methodology", 1)

    _heading(doc, "3.1  Monitoring Records", 2)
    _heading(doc, "3.1.1  Site selection criteria", 3)
    _body(doc,
          "The monitoring location was selected so that it represents the "
          "land-use pattern prescribed in the standard — industrial, "
          "commercial, residential or silence zone — and so that the "
          "measured levels are representative of background conditions at "
          "the project site.")
    _heading(doc, "3.1.2  Site conditions", 3)
    _bullets(doc, [
        "Site conditions and noise sources were recorded on a standard "
        "record sheet throughout the survey.",
        "Measurement was paused during periods of high intrusive noise "
        "unrelated to the site, where practicable.",
        "Meteorological information was used in the analysis of the noise "
        "data.",
    ])
    if getattr(campaign, "site_conditions_note", None):
        _body(doc, campaign.site_conditions_note)

    _heading(doc, "3.2  Equipment used for measurements", 2)
    meter = campaign.meter_model or "Class 1 integrating sound level meter"
    serial = f", serial number {campaign.meter_serial}" \
        if campaign.meter_serial else ""
    _body(doc,
          f"A Class 1 sound level meter was used for the survey — "
          f"{meter}{serial} — conforming to IEC 61672-1 Class 1 and fitted "
          f"with a windscreen.")

    # The instrument list comes from the equipment library, so the report
    # names the meter that was actually selected rather than whatever was
    # typed on the campaign.
    inst_rows = []
    for inst in (getattr(campaign, "instruments", None) or []):
        get = (inst.get if isinstance(inst, dict)
               else lambda k, d=None: getattr(inst, k, d))
        inst_rows.append([get("parameter") or "—",
                          get("technique") or "—",
                          get("sn") or "—",
                          get("calibration_date") or "—"])
    if inst_rows:
        _caption(doc, f"Table {tnum()} — Equipment used for the survey")
        _table(doc, ["Instrument", "Make and model", "Serial number",
                     "Calibrated"], inst_rows, [42, 62, 34, 26],
               aligns=["left", "left", "center", "center"])

    photos = [p for p in (equipment_photo_paths or [])
              if p and os.path.exists(p)]
    if photos:
        t = doc.add_table(rows=1, cols=min(2, len(photos)))
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for k, path in enumerate(photos[:2]):
            cp = _tight(t.cell(0, k).paragraphs[0], 2, 2)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pic(cp, path, 74)
        _caption(doc, f"Figure {_fn()} — Equipment used for the survey")

    _heading(doc, "3.3  Equipment measuring procedures", 2)
    height = getattr(campaign, "mic_height_m", None) or 1.5
    cal_db = getattr(campaign, "calibration_level_db", None) or 94
    _bullets(doc, [
        f"The sound level meter was mounted {height:g} m above ground "
        f"level, away from reflecting surfaces.",
        "The equipment was fixed on a tripod to prevent vibration.",
        f"The instrument was calibrated with the acoustic calibrator at "
        f"{cal_db:g} dB before the survey.",
        "Logging was started and the location attended in shifts, with key "
        "observations recorded.",
        f"On completion of the monitoring period the calibration was "
        f"checked again at {cal_db:g} dB and the stored data downloaded.",
    ])

    _heading(doc, "3.4  Calibration procedure", 2)
    calibrator = getattr(campaign, "calibrator_model", None)
    _body(doc,
          f"The meter was field-calibrated with an acoustic calibrator"
          f"{f' ({calibrator})' if calibrator else ''} at {cal_db:g} dB "
          f"immediately before and immediately after the survey. No "
          f"significant drift was observed."
          + (f" The laboratory calibration certificate is reproduced in "
             f"Appendix {_appendix['certificates']}."
             if "certificates" in _appendix else "")
          + " Data are invalidated where calibration or instrument faults "
            "affected the record.")

    _heading(doc, "3.5  Measurement procedure", 2)
    _body(doc,
          f"The A-weighted equivalent level was logged continuously at "
          f"{interval} intervals over the full survey period. Levels "
          f"outside the physically plausible range for an outdoor "
          f"measurement are flagged invalid on ingest and excluded from "
          f"every statistic in this report.")

    # Meteorology — printed only when the team recorded it, because an
    # empty weather table is worse than none at all.
    met_rows = []
    for label, value, unit in (
            ("Ambient temperature — hourly maximum",
             getattr(campaign, "met_temp_max_c", None), "°C"),
            ("Ambient temperature — hourly minimum",
             getattr(campaign, "met_temp_min_c", None), "°C"),
            ("Relative humidity — hourly maximum",
             getattr(campaign, "met_rh_max_pct", None), "%"),
            ("Relative humidity — hourly minimum",
             getattr(campaign, "met_rh_min_pct", None), "%"),
            ("Wind speed — hourly maximum",
             getattr(campaign, "met_wind_max_ms", None), "m/s"),
            ("Wind speed — hourly minimum",
             getattr(campaign, "met_wind_min_ms", None), "m/s"),
            ("Mean wind speed",
             getattr(campaign, "met_wind_mean_ms", None), "m/s")):
        if value is not None:
            met_rows.append([label, f"{value:g} {unit}"])
    prevailing = getattr(campaign, "met_wind_prevailing", None)
    if prevailing:
        met_rows.append(["Prevailing wind direction", prevailing])
    if met_rows:
        _heading(doc, "3.6  Meteorological conditions", 2)
        _body(doc,
              "Meteorological conditions recorded at the location over the "
              "monitoring period, used in the analysis of the noise data.")
        _caption(doc, f"Table {tnum()} — Monitored meteorological parameters")
        _table(doc, ["Parameter", "Result"], met_rows, [110, 54],
               aligns=["left", "center"])

    # ---- standards -------------------------------------------------------
    _heading(doc, "4  Environmental Noise Standards", 1)
    _body(doc,
          "Measured levels are assessed against the Implementing "
          "Regulations for Noise of the Environmental Law issued by Royal "
          "Decree No. (M/165), administered by the National Center for "
          "Environmental Compliance (NCEC).")
    _caption(doc, f"Table {tnum()} — Noise limits for residential and "
                  f"commercial areas")
    _table(doc, ["Category", "Description", "Daytime LAeq dB(A)",
                 "Night-time LAeq dB(A)"],
           [["A", "Sensitive zones — hospitals, schools", "50", "40"],
            ["B", "Residential areas", "55", "45"],
            ["C", "Mixed residential and commercial", "60", "50"],
            ["D", "Commercial and business districts", "65", "55"]],
           [18, 74, 36, 36], aligns=["center", "left", "center", "center"])
    _caption(doc, f"Table {tnum()} — Noise limits for roadsides")
    _table(doc, ["Zone", "Daytime LAeq dB(A)", "Night-time LAeq dB(A)"],
           [["Main roads and highways", "70", "65"]],
           [92, 36, 36], aligns=["left", "center", "center"])
    _caption(doc, f"Table {tnum()} — Noise limits for industrial zones")
    _table(doc, ["Zone", "Daytime LAeq dB(A)", "Night-time LAeq dB(A)"],
           [["Industrial zones", "70", "65"]],
           [92, 36, 36], aligns=["left", "center", "center"])
    corrections_table = tnum()
    _caption(doc, f"Table {corrections_table} — Corrections applied at "
                  f"construction work sites")
    _table(doc, ["Period of construction activities",
                 "Correction to allowed level dB(A)"],
           [[p, f"+{c}"] for p, c in CONSTRUCTION_CORRECTIONS],
           [110, 54], aligns=["left", "center"])

    # ---- results ---------------------------------------------------------
    _heading(doc, "5  Results", 1)
    _body(doc,
          f"Statistical parameters were derived from the validated "
          f"{interval} record by energy averaging. The full record, the "
          f"hourly profile and the statistical distribution are presented "
          f"in the figures that follow.")
    _caption(doc, f"Table {tnum()} — Summary of measured noise levels")
    _table(doc, ["Site", "LAeq, T", "L Day", "L Night", "LA10", "LA50",
                 "LA90", "Lmax", "Lmin"],
           [["N1", _f(summary.laeq_t), _f(summary.l_day),
             _f(summary.l_night), _f(summary.la10), _f(summary.la50),
             _f(summary.la90), _f(summary.lmax), _f(summary.lmin)]],
           [20, 18, 18, 18, 18, 18, 18, 18, 18])
    _muted(doc,
           f"Values in dB(A). L Day and L Night are energy averages over "
           f"{summary.day_start_hour:02d}:00–{summary.day_end_hour:02d}:00 "
           f"and the complementary night period respectively.")
    _caption(doc, f"Table {tnum()} — Data capture ({interval} intervals)")
    _table(doc, ["Expected intervals", "Recorded", "Valid", "Invalid",
                 "Capture %"],
           [[f"{summary.expected_records:,}", f"{summary.total_records:,}",
             f"{summary.valid_records:,}", f"{summary.invalid_records:,}",
             f"{summary.data_capture_pct}"]],
           [36, 32, 32, 32, 32])

    for key, cap_text in (("hourly", "Hourly LAeq over the survey with the "
                                     "applicable day and night limits"),
                          ("daynight", "Day and night levels against the "
                                       "applicable NCEC limit"),
                          ("cats", "Measured levels against the four NCEC "
                                   "land-use categories"),
                          ("trace", f"Sound level record ({interval} "
                                    f"resolution) with the LA90–LA10 "
                                    f"statistical envelope"),
                          ("dist", "Exceedance distribution of measured "
                                   "levels with statistical percentiles")):
        path = figs.get(key)
        if not path or not os.path.exists(path):
            continue
        p = _tight(doc.add_paragraph(), 6, 0)
        pic(p, path, 164)
        _caption(doc, f"Figure {_fn()} — {cap_text}")

    # ---- assessment ------------------------------------------------------
    _heading(doc, "6  Discussion and Compliance", 1)
    limit = NOISE_LIMITS.get(campaign.noise_category)
    if verdicts and limit:
        _body(doc,
              f"The applicable land-use category for this location is "
              f"{limit.label}. The measured day-time and night-time levels "
              f"are assessed against its limits below.")
        for v in verdicts:
            if v.status == "ok":
                _callout(doc, v.text, GREEN_HEX, "EEF6F1", "\u2713")
            else:
                _callout(doc, v.text, AMBER_HEX, "FDF3E6", "\u26a0")
        if campaign.noise_category == "construction" or any(
                v.status == "over" for v in verdicts):
            _body(doc,
                  f"Where activities at the location constitute "
                  f"construction "
                  f"work, the corrections of Table {corrections_table} "
                  f"apply to the allowed levels according to the duration "
                  f"of the activity.")
    else:
        _body(doc,
              "The most appropriate national standard category for ambient "
              "noise at this location is to be determined by the "
              "consultant. The measured levels are stated against every "
              "category in Figure form above; no compliance judgement is "
              "made in this report.")

    _heading(doc, "7  Conclusion", 1)
    _body(doc,
          f"Noise monitoring was conducted at one location for "
          f"{dur.total_seconds() / 3600:.0f} hours from "
          f"{_fmt_dt(campaign.monitoring_start)}. The summarised data were "
          f"derived from the validated logged record to provide "
          f"energy-averaged levels over the day-time and night-time "
          f"periods, together with the statistical parameters of Table 7.")

    _heading(doc, "8  References", 1)
    _body(doc, "[1]  Implementing Regulations for Noise of the "
               "Environmental Law, National Center for Environmental "
               "Compliance (NCEC), Kingdom of Saudi Arabia.")
    _body(doc, "[2]  IEC 61672-1, Electroacoustics — Sound level meters — "
               "Part 1: Specifications.")

    _heading(doc, "9  Glossary", 1)
    _table(doc, ["Term", "Definition"],
           [["dB(A)", "A-weighted sound pressure level, weighted to the "
                      "response of the human ear."],
            ["LAeq, T", "Equivalent continuous A-weighted level over "
                        "period T — the energy average of the record."],
            ["L Day / L Night", "LAeq over the day-time and night-time "
                                "periods defined by the regulation."],
            ["LA10 / LA50 / LA90", "Levels exceeded 10%, 50% and 90% of "
                                   "the time respectively."],
            ["Lmax / Lmin", "Highest and lowest logged interval levels."]],
           [40, 124], aligns=["left", "left"])

    # ---- appendices ------------------------------------------------------
    if calibration_items:
        doc.add_page_break()
        _heading(doc, f"Appendix {_appendix['certificates']} — "
                      f"Calibration certificates", 1)
        _body(doc,
              "The laboratory calibration certificates covering the "
              "instrument used for this survey are reproduced below.")
        for i, c in enumerate(calibration_items):
            path = c.get("path")
            if not (path and os.path.exists(path)):
                continue
            if i:
                doc.add_page_break()
            title = (c.get("title") or "Calibration certificate").strip()
            p = _tight(doc.add_paragraph(), 2, 4)
            r = p.add_run(title)
            r.font.size = Pt(9.5)
            r.font.bold = True
            r.font.color.rgb = NAVY
            p = _tight(doc.add_paragraph(), 0, 0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pic(p, path, 165)
    if license_image_paths:
        doc.add_page_break()
        _heading(doc, f"Appendix {_appendix['licence']} — Environmental "
                      f"license for the institution", 1)
        for i, path in enumerate(license_image_paths):
            if not (path and os.path.exists(path)):
                continue
            if i:
                doc.add_page_break()
            p = _tight(doc.add_paragraph(), 2, 0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pic(p, path, 165)

    # ---- end marker ------------------------------------------------------
    doc.add_page_break()
    for _ in range(3):
        _tight(doc.add_paragraph(), 0, 0)
    p = _tight(doc.add_paragraph(), 0, 2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("END OF REPORT")
    r.font.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = NAVY
    p = _tight(doc.add_paragraph(), 0, 0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("  ·  ".join(
        x for x in (campaign.report_number,
                    f"Rev {campaign.revision or '00'}",
                    campaign.project_name) if x))
    r.font.size = Pt(9)
    r.font.color.rgb = MUT

    os.makedirs(os.path.dirname(os.path.abspath(out_path)), exist_ok=True)
    doc.save(out_path)
    return out_path
